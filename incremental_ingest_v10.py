import os
import json
import time
import uuid
import sqlite3
import hashlib
import re
import argparse
import difflib
from datetime import datetime
import pdfplumber
from docx import Document
import google.generativeai as genai
from neo4j import GraphDatabase
from tqdm import tqdm
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed

from connectors.gdrive_api import GDriveConnector
from googleapiclient.http import MediaFileUpload
from openai import OpenAI
from connectors.pinecone_api import PineconeClient
from batch_pinecone_sync import chunk_text

# Locks
db_lock = threading.Lock()
neo_lock = threading.Lock()
pinecone_lock = threading.Lock()
gdrive_lock = threading.Lock()

# Target Folders
TARGET_DIRS = [
    r"C:\Users\cazam\Downloads\02_resume 전처리",
    r"C:\Users\cazam\Downloads\02_resume_converted_v8"
]
DB_PATH = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\candidates.db"

with open(r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\secrets.json", "r", encoding="utf-8") as f:
    secrets = json.load(f)

# Clients
genai.configure(api_key=secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel("gemini-2.5-flash")
n_uri = secrets.get("NEO4J_URI", "neo4j://127.0.0.1:7687")
n_user = secrets.get("NEO4J_USERNAME", "neo4j")
n_pw = secrets.get("NEO4J_PASSWORD", "toss1234")
driver = GraphDatabase.driver(n_uri, auth=(n_user, n_pw))

gdrive = GDriveConnector()
folder_id = secrets.get("GOOGLE_DRIVE_FOLDER_ID")

openai_client = OpenAI(api_key=secrets.get("OPENAI_API_KEY", ""))
pc_host = secrets.get("PINECONE_HOST", "").rstrip("/")
if not pc_host.startswith("https://"): pc_host = f"https://{pc_host}"
pinecone_client = PineconeClient(secrets.get("PINECONE_API_KEY", ""), pc_host)

def init_db():
    # Ensure new columns exist
    with sqlite3.connect(DB_PATH) as conn:
        c = conn.cursor()
        for col, ctype in [("birth_year", "INTEGER"), ("total_years", "REAL"), 
                           ("sector", "TEXT"), ("education_json", "TEXT"), 
                           ("profile_summary", "TEXT"), ("current_company", "TEXT")]:
            try:
                c.execute(f"ALTER TABLE candidates ADD COLUMN {col} {ctype}")
            except:
                pass

MEGA_PROMPT = """이력서 텍스트를 분석해서 아래 구조를 JSON Object 형식으로 출력해. 반드시 코드블럭 없이 JSON만 반환할 것.

{{
  "name_kr": "본문에서 순수 본명만 추출 (직무 꼬리표 붙이지 말 것). 못 찾으면 null",
  "phone": "010-XXXX-XXXX 형태 숫자만. 못 찾으면 null",
  "email": "이메일주소. 못 찾으면 null",
  "birth_year": 출생연도 숫자 4자리 (예: 1990). 못 찾으면 null,
  "summary": "주요 경력 2줄 요약 (전화번호/이메일/주소/생년월일 등 개인정보 절대 포함 금지!!!)",
  "sector": "가장 관련 깊은 도메인/산업분야 분류 (예: SW, Finance, Marketing 등. 미분류 금지)",
  "education_json": [
    {{"school": "대학교 이상 학교명", "major": "전공", "degree": "학사/석사 등", "year": "졸업년도"}}
  ],
  "careers_json": [
    {{
      "company": "회사명",
      "title": "직책/부서",
      "start_date": "YYYY.MM",
      "end_date": "YYYY.MM (또는 현재)"
    }}
  ],
  "neo4j_edges": [
    {{
      "action": "BUILT|DESIGNED|MANAGED|ANALYZED|LAUNCHED|NEGOTIATED|GREW|SUPPORTED 중 택1",
      "skill": "아래 지시된 Skill 목록 중 택1",
      "confidence": 0.0 ~ 1.0,
      "evidence_span": "해당 판단을 내리게 된 본문 문구 복사"
    }}
  ]
}}

[Skill 목록]:
Payment_and_Settlement_System, Service_Planning, Product_Manager, Data_Pipeline_Construction, Data_Engineering, Data_Analysis,
Backend, Frontend, Machine_Learning, MLOps, DevOps, Financial_Accounting, Corporate_Strategic_Planning, 사업개발_BD, 퍼포먼스마케팅,
Treasury_Management, FX_Dealing, Corporate_Funding, IPO_Preparation_and_Execution, Recruiting_and_Talent_Acquisition,
Organizational_Development, B2B영업, 물류_Logistics, Backend_Python, Backend_Java, Backend_Go, Backend_Node,
Kubernetes, Infrastructure_and_Cloud, 보안_Security, FinTech, Natural_Language_Processing, 컴퓨터비전_CV, 추천시스템,
Deep_Learning, Corporate_Legal_Counsel, Intellectual_Property, Legal_Compliance, Contract_Management, Litigation

이력서 본문:
{text}
"""

def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    text = ""
    try:
        import fitz
        if ext == "pdf":
            with fitz.open(filepath) as doc:
                for page in doc: text += page.get_text()
        elif ext in ("docx", "doc"):
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            for t in doc.tables:
                for r in t.rows:
                    for ce in r.cells:
                        if ce.text: text += ce.text + " "
    except Exception as e:
        print(f"Extraction error for {filepath}: {e}")
    return text.strip()

def calculate_career_stats(careers):
    if not careers: return "미상", 0.0
    total_months = 0
    latest_end_dt = datetime(1900, 1, 1)
    curr_comp = "미상"
    
    def parse_dt(d_str):
        if not d_str or '현재' in d_str or '재직' in d_str or 'ing' in d_str.lower(): return datetime.now()
        m = re.findall(r'(\d{4})[^\d]*(\d{1,2})', d_str)
        if m: return datetime(int(m[0][0]), int(m[0][1]), 1)
        m2 = re.findall(r'(\d{4})', d_str)
        if m2: return datetime(int(m2[0]), 1, 1)
        return datetime.now()

    for c in careers:
        st = parse_dt(c.get('start_date', ''))
        ed = parse_dt(c.get('end_date', ''))
        if ed < st: ed = st
        months = (ed.year - st.year) * 12 + ed.month - st.month
        total_months += max(months, 0)
        
        if ed >= latest_end_dt:
            latest_end_dt = ed
            if c.get('company'): curr_comp = c.get('company')
            
    total_y = round(total_months / 12.0, 1)
    return curr_comp, total_y

def extract_fallback_name(filename):
    m = re.search(r'[가-힣]{2,4}', filename.replace("이력서", "").replace("포트폴리오", "").replace("개발자", ""))
    return m.group(0) if m else "미상"

def process_file(filepath):
    filename = os.path.basename(filepath)
    text = extract_text(filepath)
    if len(text) < 50:
        return False, "Not enough text"
        
    doc_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    
    # 1. MD5 Hash Duplicate Check
    with db_lock:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM candidates WHERE document_hash=?', (doc_hash,))
        count = c.fetchone()[0]
        if count > 0:
            conn.close()
            return False, "기존 파일 스킵 (MD5)"
            
        # 2. Name + Text similarity check
        fallback_name = extract_fallback_name(filename)
        if fallback_name and fallback_name != "미상":
            c.execute('SELECT raw_text FROM candidates WHERE name_kr LIKE ? AND is_duplicate=0', (f"%{fallback_name}%",))
            matches = c.fetchall()
            for db_row in matches:
                db_text = db_row[0]
                if not db_text: continue
                # Compare first 400 chars
                similarity = difflib.SequenceMatcher(None, text[:400], db_text[:400]).ratio()
                if similarity >= 0.7:
                    conn.close()
                    return False, f"기존 파일 스킵 (이름 매칭 & 텍스트 유사도 {int(similarity*100)}%)"

        conn.close()

    # 2. Upload to Google Drive FIRST
    drive_link = None
    with gdrive_lock:
        try:
            drive_res = gdrive.service.files().list(q=f"name='{filename}' and '{folder_id}' in parents and trashed=false", fields="files(id, webViewLink)").execute()
            if drive_res.get('files'):
                drive_link = drive_res['files'][0]['webViewLink']
            else:
                ext = filename.lower().split('.')[-1]
                mt = 'application/pdf' if ext == 'pdf' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                media = MediaFileUpload(filepath, mimetype=mt, resumable=True)
                uploaded = gdrive.service.files().create(body={'name': filename, 'parents': [folder_id]}, media_body=media, fields='webViewLink').execute()
                drive_link = uploaded.get('webViewLink')
        except Exception as e:
            print(f"Drive Upload Error: {e}")

    # 3. Gemini Parsing
    parsed = None
    for attempt in range(3):
        try:
            prompt = MEGA_PROMPT.replace("{text}", f"[파일명: {filename}]\n\n" + text[:6000])
            res = model.generate_content(prompt)
            raw = res.text.replace("```json", "").replace("```", "").strip()
            parsed = json.loads(raw)
            break
        except Exception as e:
            time.sleep(2)
            
    if not parsed:
        return False, "Gemini Parsing Failed"

    name_kr = parsed.get("name_kr")
    if not name_kr: name_kr = extract_fallback_name(filename)

    email = parsed.get("email", "")
    phone = parsed.get("phone", "")
    birth_year = parsed.get("birth_year", 0)
    sector = parsed.get("sector", "미분류")
    summary = parsed.get("summary", "")
    careers = parsed.get("careers_json", [])
    edu = parsed.get("education_json", [])
    
    current_company, total_years = calculate_career_stats(careers)

    # 4. Duplicate Check (Phone/Email)
    is_duplicate = 0
    with db_lock:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        
        # Strip phone numbers for comparison
        clean_ext_phone = phone.replace("-","").strip() if phone else ""
        clean_ext_email = email.lower().strip() if email else ""
        
        if clean_ext_phone or clean_ext_email:
            for row in c.execute("SELECT id, email, phone FROM candidates"):
                d_email = row[1].lower().strip() if row[1] else ""
                d_phone = row[2].replace("-","").strip() if row[2] else ""
                
                if (clean_ext_email and d_email == clean_ext_email) or \
                   (clean_ext_phone and d_phone == clean_ext_phone):
                    is_duplicate = 1
                    print(f"  [Duplicate Found] {name_kr}: Match with existing -> Email: {clean_ext_email}, Phone: {clean_ext_phone}")
                    break
                    
        c_id = str(uuid.uuid4())
        c.execute('''
            INSERT INTO candidates (
                id, name_kr, email, phone, birth_year, sector, profile_summary, total_years, current_company,
                google_drive_url, raw_text, document_hash, is_duplicate, is_parsed, is_neo4j_synced, is_pinecone_synced,
                careers_json, education_json, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1, 1, 1, ?, ?, datetime('now'), datetime('now'))
        ''', (c_id, name_kr, email, phone, birth_year, sector, summary, total_years, current_company,
              drive_link, text, doc_hash, is_duplicate, json.dumps(careers, ensure_ascii=False), json.dumps(edu, ensure_ascii=False)))
        conn.commit()
        conn.close()

    # If duplicate, we still inserted it with is_duplicate=1. We skip Graph/Vector DB to prevent junk.
    if is_duplicate:
        return True, f"{name_kr} (Duplicate Saved)"

    # 5. Neo4j Sync
    with neo_lock:
        with driver.session() as session:
            session.run("""
                MERGE (c:Candidate {id: $id})
                SET c.name = $name_kr, c.phone = $phone, c.email = $email,
                    c.current_company = $current_company, c.profile_summary = $summary,
                    c.total_years = $total_years, c.sector = $sector
            """, id=c_id, name_kr=name_kr, phone=phone, email=email, 
                 current_company=current_company, summary=summary, total_years=total_years, sector=sector)
                 
            for edge in parsed.get("neo4j_edges", []):
                act, skill, conf, ev = edge.get("action", ""), edge.get("skill", ""), float(edge.get("confidence", 0.5)), edge.get("evidence_span", "")
                if act and skill:
                    session.run(f"""
                        MERGE (c:Candidate {{id: $id}})
                        MERGE (s:Skill {{name: $skill}})
                        MERGE (c)-[r:{act}]->(s)
                        SET r.confidence = $conf, r.evidence_span = $ev, r.source = 'v10_ingest'
                    """, id=c_id, skill=skill, conf=conf, ev=ev)

    # 6. Pinecone Embeddings
    with pinecone_lock:
        try:
            chunks = chunk_text(text)
            if chunks:
                response = openai_client.embeddings.create(model="text-embedding-3-small", input=chunks)
                vectors = []
                for i, emb in enumerate(response.data):
                    vectors.append({
                        "id": f"{c_id}_chunk_{i}",
                        "values": emb.embedding,
                        "metadata": {"candidate_id": c_id, "chunk_index": i}
                    })
                pinecone_client.upsert(vectors, namespace="resume_vectors")
        except Exception as e:
            print(f"Pinecone error: {e}")

    return True, f"{name_kr} (total: {total_years} yrs, {current_company})"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--test', action='store_true', help='Run in test mode (max 3 files)')
    args = parser.parse_args()

    init_db()

    files = []
    for d in TARGET_DIRS:
        if os.path.exists(d):
            files.extend([os.path.join(d, f) for f in os.listdir(d) 
                         if f.endswith(('.pdf', '.docx', '.doc')) and not f.startswith('~$')])
            
    print(f"Total target files found: {len(files)}")
    if args.test:
        files = files[:20]
        print("TEST MODE: Limiting to 20 files.")
        
    success = 0
    skipped = 0
    duplicate_saved = 0
    total_processed = 0
    
    with ThreadPoolExecutor(max_workers=3) as executor:
        futures = {executor.submit(process_file, fp): os.path.basename(fp) for fp in files}
        for future in tqdm(as_completed(futures), total=len(files)):
            filename = futures[future]
            try:
                ok, res_str = future.result()
                if ok:
                    if "Duplicate Saved" in res_str:
                        print(f"\n[DUPLICATE] {filename} -> {res_str}")
                        duplicate_saved += 1
                    else:
                        print(f"\n[OK] {filename} -> {res_str}")
                        success += 1
                else:
                    print(f"\n[SKIP/ERR] {filename} -> {res_str}")
                    skipped += 1
            except Exception as e:
                skipped += 1
                print(f"\n[ERROR] {filename}: {e}")
                
            total_processed += 1
            if total_processed % 100 == 0:
                print(f"\n[Progress {total_processed}/{len(files)}] Success: {success} | Skipped/Err: {skipped} | Duplicates: {duplicate_saved}")
            
    print(f"\n[V10 Complete] Success(New): {success} | Skipped: {skipped} | Duplicate(3rd step): {duplicate_saved}")

if __name__ == "__main__":
    main()
