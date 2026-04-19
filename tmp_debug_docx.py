import os
from docx import Document

def extract_full_docx(filepath):
    try:
        doc = Document(filepath)
        full_text = []
        
        # 1. Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # 2. Tables (CRITICAL for this resume)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except Exception as e:
        return f"ERROR: {e}"

if __name__ == "__main__":
    p = r"C:\Users\cazam\Downloads\02_resume 전처리\[NC] 현종민(IR)부문_86.docx"
    content = extract_full_docx(p)
    print(f"TOTAL LENGTH: {len(content)}")
    print("--- CONTENT START ---")
    print(content[:5000])
    print("--- CONTENT END ---")
