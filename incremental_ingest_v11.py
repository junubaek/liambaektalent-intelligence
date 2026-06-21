"""
incremental_ingest_v11.py
- GPT-4.1-mini 구조화 추출
- sector 복수 선택 (세미콜론 구분)
- current_title 컬럼 추가
- 대기업/스타트업 판별
- 총경력 산정 (동시재직 중복 제거)
- Google Drive 업로드 + 링크
- 변환본 폴더 우선, 원본 폴더 보완
- MD5 기반 중복 제거
"""
import os, json, time, uuid, sqlite3, hashlib, re, difflib
import threading
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz
from docx import Document
from neo4j import GraphDatabase
from tqdm import tqdm
from openai import OpenAI
from googleapiclient.http import MediaFileUpload

from connectors.gdrive_api import GDriveConnector
from connectors.pinecone_api import PineconeClient
from batch_pinecone_sync import chunk_text
from ontology_graph import CANONICAL_MAP

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
DB_PATH = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\candidates.db"
TARGET_DIRS = [
    r"C:\Users\cazam\Downloads\02_resume_converted_v8",  # 변환본 우선
    r"C:\Users\cazam\Downloads\02_resume 전처리",          # 원본 보완
]

with open(r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\secrets.json", encoding="utf-8") as _f:
    _s = json.load(_f)

GPT_MODEL     = "gpt-4.1-mini"
openai_client = OpenAI(api_key=_s["OPENAI_API_KEY"])
neo4j_driver  = GraphDatabase.driver(_s["NEO4J_URI"], auth=(_s["NEO4J_USERNAME"], _s["NEO4J_PASSWORD"]))
gdrive        = GDriveConnector()
folder_id     = _s.get("GOOGLE_DRIVE_FOLDER_ID", "")

_pc_host = _s.get("PINECONE_HOST", "").rstrip("/")
if not _pc_host.startswith("https://"):
    _pc_host = f"https://{_pc_host}"
pinecone_client = PineconeClient(_s["PINECONE_API_KEY"], _pc_host)

db_lock       = threading.Lock()
neo_lock      = threading.Lock()
pinecone_lock = threading.Lock()
gdrive_lock   = threading.Lock()

# ---------------------------------------------------------------------------
# Sector 목록 (복수 가능, 세미콜론 구분)
# ---------------------------------------------------------------------------
VALID_SECTORS = {
    "Eng_SW", "Eng_AI", "Eng_Data", "Eng_Semi", "Eng_Embedded", "Eng_HW",
    "Finance", "Strategy", "Sales", "Marketing", "HR", "Operations",
    "Product", "Design", "Legal", "Healthcare", "C_Level",
}

def validate_sector(raw: str) -> str:
    """세미콜론으로 구분된 sector 문자열 검증 및 정규화"""
    if not raw:
        return "Eng_SW"
    parts = [p.strip() for p in raw.split(";") if p.strip()]
    valid = [p for p in parts if p in VALID_SECTORS]
    return ";".join(valid) if valid else "Eng_SW"

# ---------------------------------------------------------------------------
# 대기업 / 스타트업 판별
# ---------------------------------------------------------------------------
BIGCORP = {
    "삼성", "samsung", "lg", "엘지", "sk", "현대", "hyundai", "기아",
    "롯데", "한화", "cj", "네이버", "naver", "카카오", "kakao",
    "쿠팡", "coupang", "배달의민족", "우아한형제", "토스", "toss",
    "당근", "라인", "line", "크래프톤", "넥슨", "엔씨소프트", "넷마블",
    "kt", "포스코", "posco", "두산", "gs", "구글", "google",
    "메타", "meta", "아마존", "amazon", "마이크로소프트", "microsoft",
    "apple", "애플", "hybe", "하이브", "sm", "jyp", "yg",
    "올리브영", "아모레퍼시픽", "오리온", "농심", "맥킨지", "bcg",
    '코웨이', 'coway', '교원', '대원미디어', '파이낸셜데이타',
    '현대백화점', '신세계', '이마트', '롯데백화점', '갤러리아',
    '한국타이어', '금호', '효성', '코오롱', '한진', '대한항공',
    '아시아나', '하나금융', '신한', '우리은행', '기업은행', 'kb',
    '삼성화재', '현대해상', '교보생명', '한화생명', '미래에셋',
    '카카오페이', '카카오뱅크', '토스뱅크', '케이뱅크',
    '직방', '야놀자', '여기어때', '마켓컬리', '오아시스',
    '무신사', '에이블리', '지그재그', '브랜디',
    "골드만", "goldman", "jp morgan", "삼정kpmg", "딜로이트", "pwc",
    "한국전력", "kepco", "인터파크", "위메프", "11번가", "지마켓",
}

STARTUP_SIGNALS = {
    "스타트업", "벤처", "series a", "series b", "seed", "시드",
    "pre-ipo", "엑셀러레이터", "인큐베이터",
}

def detect_company_type(company: str):
    if not company:
        return 0, 0
    name = company.lower()
    has_big     = int(any(k in name for k in BIGCORP))
    has_startup = int(any(k in name for k in STARTUP_SIGNALS))
    return has_big, has_startup

# ---------------------------------------------------------------------------
# DB 초기화
# ---------------------------------------------------------------------------
def init_db():
    with sqlite3.connect(DB_PATH) as conn:
        c = conn.cursor()
        for col, ctype in [
            ("current_title",    "TEXT"),
            ("has_big_company",  "INTEGER DEFAULT 0"),
            ("has_startup",      "INTEGER DEFAULT 0"),
            ("company_timeline", "TEXT"),
        ]:
            try:
                c.execute(f"ALTER TABLE candidates ADD COLUMN {col} {ctype}")
            except Exception:
                pass

# ---------------------------------------------------------------------------
# GPT 프롬프트
# ---------------------------------------------------------------------------
SECTOR_LIST = ", ".join(sorted(VALID_SECTORS))

MEGA_PROMPT = f"""이력서 텍스트를 분석해서 아래 JSON을 반환하라.
코드블럭 없이 JSON만 반환. 이력서 본문에 없는 내용은 절대 만들지 말 것.

{{
  "name_kr": "본명만. 직무명/회사명 붙이지 말 것. 없으면 null",
  "phone": "010-XXXX-XXXX 형태. 없으면 null",
  "email": "이메일. 없으면 null",
  "birth_year": 출생연도 4자리 숫자 또는 null,
  "current_title": "현재 직책/포지션. 예: Senior Engineer, 팀장, Director",
  "summary": "핵심 역량과 경력 2-3줄 요약. 개인정보(전화/이메일/주소) 절대 포함 금지",
  "sector": "아래 목록에서 선택. 복수 가능시 세미콜론으로 구분. 예: Eng_SW;Eng_AI\\n가능한 값: {SECTOR_LIST}",
  "education_json": [
    {{"school": "대학교 이상", "major": "전공", "degree": "학사/석사/박사", "year": "졸업연도"}}
  ],
  "careers_json": [
    {{
      "company": "회사명",
      "title": "직책",
      "start_date": "YYYY.MM",
      "end_date": "YYYY.MM 또는 현재",
      "description": "주요 업무/성과 1-2줄"
    }}
  ],
  "neo4j_edges": [
    {{
      "action": "BUILT|DESIGNED|MANAGED|ANALYZED|LAUNCHED|NEGOTIATED|GREW|SUPPORTED 중 택1",
      "skill": "구체적 스킬명 (영문 권장)",
      "confidence": 0.0~1.0,
      "evidence_span": "근거 문구 (원문 그대로)"
    }}
  ]
}}

sector 선택 가이드:
- 개발자/엔지니어: Eng_SW (백엔드/프론트/DevOps), Eng_AI (ML/DL), Eng_Data (데이터), Eng_Semi (반도체), Eng_Embedded (펌웨어/드라이버), Eng_HW (하드웨어/전기)
- 비즈니스: Finance (재무/회계/FP&A/급여제외), HR (인사/채용/급여/총무), Strategy (전략/기획/M&A/BD), Sales (영업), Marketing (마케팅/PR), Operations (운영/물류/전력계통)
- 기타: Product (PM/PO), Design (UI/UX), Legal (법무), Healthcare (의료/바이오), C_Level (CEO/COO/CFO/CTO/임원)
- 복수 예시: 반도체 SW → Eng_Semi;Eng_SW | CTO → C_Level;Eng_SW | MLOps → Eng_AI;Eng_SW

이력서:
{{text}}
"""

# ---------------------------------------------------------------------------
# 헬퍼
# ---------------------------------------------------------------------------
def extract_text(filepath: str) -> str:
    ext = filepath.rsplit(".", 1)[-1].lower()
    text = ""
    try:
        if ext == "pdf":
            with fitz.open(filepath) as doc:
                text = "".join(p.get_text() for p in doc)
        elif ext in ("docx", "doc"):
            doc = Document(filepath)
            parts = [p.text for p in doc.paragraphs]
            for tbl in doc.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            text = "\n".join(parts)
        elif ext == "txt":
            with open(filepath, encoding="utf-8", errors="ignore") as f:
                text = f.read()
    except Exception as e:
        print(f"[extract_text 오류] {os.path.basename(filepath)}: {e}")
    return text.strip()

def extract_fallback_name(filename: str) -> str:
    clean = re.sub(r"이력서|포트폴리오|개발자|resume|cv", "", filename, flags=re.IGNORECASE)
    m = re.search(r"[\uAC00-\uD7A3]{2,4}", clean)
    return m.group(0) if m else "미상"

def calculate_career_stats(careers: list):
    if not careers:
        return "미상", 0.0, "[]"

    def parse_dt(d):
        if not d:
            return datetime.now()
        if any(k in str(d) for k in ["현재", "재직", "present", "ing"]):
            return datetime.now()
        m = re.findall(r"(\d{4})[^\d]*(\d{1,2})", str(d))
        if m:
            return datetime(int(m[0][0]), max(1, min(12, int(m[0][1]))), 1)
        m2 = re.findall(r"(\d{4})", str(d))
        if m2:
            return datetime(int(m2[0]), 1, 1)
        return datetime.now()

    timeline = []
    periods  = []
    latest   = datetime(1900, 1, 1)
    current_co = "미상"

    for c in careers:
        st = parse_dt(c.get("start_date"))
        ed = parse_dt(c.get("end_date"))
        if ed < st:
            ed = st
        months = max(0, (ed.year - st.year) * 12 + ed.month - st.month)
        periods.append((st, ed))
        timeline.append({
            "company":  c.get("company", ""),
            "title":    c.get("title", ""),
            "start":    c.get("start_date", ""),
            "end":      c.get("end_date", ""),
            "months":   months,
            "desc":     c.get("description", ""),
        })
        if ed >= latest:
            latest = ed
            if c.get("company"):
                current_co = c["company"]

    # 동시재직 중복 제거
    periods.sort()
    merged = []
    for st, ed in periods:
        if merged and st <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], ed))
        else:
            merged.append([st, ed])

    total_months = sum(
        max(0, (e.year - s.year) * 12 + e.month - s.month)
        for s, e in merged
    )
    return current_co, round(total_months / 12.0, 1), json.dumps(timeline, ensure_ascii=False)

def normalize_skill(skill: str) -> str:
    if skill in CANONICAL_MAP:
        return CANONICAL_MAP[skill]
    low = skill.lower()
    for k, v in CANONICAL_MAP.items():
        if k.lower() == low:
            return v
    return skill

def build_embedding_text(c: dict) -> str:
    parts = [
        c.get("name_kr", ""),
        c.get("current_title", ""),
        c.get("sector", ""),
        c.get("current_company", ""),
        c.get("profile_summary", ""),
        (c.get("raw_text", "") or "")[:1500],
    ]
    return " ".join(p for p in parts if p)

# ---------------------------------------------------------------------------
# 핵심 파이프라인
# ---------------------------------------------------------------------------
def process_file(filepath: str):
    if os.path.getsize(filepath) < 1024:
        return False, "파일 너무 작음"

    filename = os.path.basename(filepath)
    text = extract_text(filepath)
    if len(text) < 50:
        return False, "텍스트 부족"

    doc_hash = hashlib.md5(text.encode("utf-8")).hexdigest()

    # MD5 중복 체크
    with db_lock:
        conn = sqlite3.connect(DB_PATH)
        c    = conn.cursor()
        c.execute("SELECT COUNT(*) FROM candidates WHERE document_hash=?", (doc_hash,))
        if c.fetchone()[0] > 0:
            conn.close()
            return False, "MD5 중복"

        # 이름+유사도 중복 체크 (임계값 0.85)
        fallback = extract_fallback_name(filename)
        if fallback != "미상":
            c.execute(
                "SELECT raw_text FROM candidates WHERE name_kr LIKE ? AND is_duplicate=0",
                (f"%{fallback}%",)
            )
            for (db_text,) in c.fetchall():
                if not db_text:
                    continue
                sim = difflib.SequenceMatcher(None, text[:400], db_text[:400]).ratio()
                if sim >= 0.85:
                    conn.close()
                    return False, f"유사도 중복 ({int(sim*100)}%)"
        conn.close()

    # Google Drive 업로드
    drive_link = None
    if folder_id:
        with gdrive_lock:
            try:
                res = gdrive.service.files().list(
                    q=f"name='{filename}' and '{folder_id}' in parents and trashed=false",
                    fields="files(id,webViewLink)"
                ).execute()
                if res.get("files"):
                    drive_link = res["files"][0]["webViewLink"]
                else:
                    ext = filename.lower().rsplit(".", 1)[-1]
                    mt  = ("application/pdf" if ext == "pdf"
                           else "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                    media    = MediaFileUpload(filepath, mimetype=mt, resumable=True)
                    uploaded = gdrive.service.files().create(
                        body={"name": filename, "parents": [folder_id]},
                        media_body=media, fields="webViewLink"
                    ).execute()
                    drive_link = uploaded.get("webViewLink")
            except Exception as e:
                print(f"[Drive 오류] {filename}: {e}")

    # GPT 추출
    parsed = None
    for attempt in range(3):
        try:
            prompt = MEGA_PROMPT.replace("{text}", f"[파일명: {filename}]\n\n{text[:6000]}")
            resp   = openai_client.chat.completions.create(
                model=GPT_MODEL,
                messages=[
                    {"role": "system", "content": "Resume analyst. Return JSON only."},
                    {"role": "user",   "content": prompt},
                ],
                temperature=0.0, max_tokens=2500,
            )
            raw_resp = resp.choices[0].message.content.strip()
            raw_resp = raw_resp.replace("```json", "").replace("```", "").strip()
            parsed   = json.loads(raw_resp)
            break
        except Exception as e:
            print(f"[GPT 오류 {attempt+1}] {filename}: {e}")
            time.sleep(2)

    if not parsed:
        return False, "GPT 파싱 실패"

    # 필드 추출
    name_kr       = parsed.get("name_kr") or fallback
    email         = parsed.get("email") or ""
    phone         = parsed.get("phone") or ""
    birth_year    = parsed.get("birth_year") or 0
    current_title = parsed.get("current_title") or ""
    summary       = parsed.get("summary") or ""
    sector        = validate_sector(parsed.get("sector", ""))
    careers       = parsed.get("careers_json", [])
    edu           = parsed.get("education_json", [])

    current_company, total_years, company_timeline = calculate_career_stats(careers)
    has_big, has_startup = detect_company_type(current_company)

    # 연락처 기반 중복 체크
    is_duplicate = 0
    with db_lock:
        conn = sqlite3.connect(DB_PATH)
        c    = conn.cursor()
        clean_email = email.lower().strip()
        clean_phone = phone.replace("-", "").strip()
        clean_name  = name_kr.strip()
        clean_co    = current_company.strip()

        for row in c.execute(
            "SELECT email, phone, name_kr, current_company FROM candidates WHERE is_duplicate=0"
        ):
            d_email = (row[0] or "").lower().strip()
            d_phone = (row[1] or "").replace("-", "").strip()
            d_name  = (row[2] or "").strip()
            d_co    = (row[3] or "").strip()

            if clean_email and d_email == clean_email:
                is_duplicate = 1; break
            if clean_phone and d_phone == clean_phone:
                is_duplicate = 1; break
            if clean_name and clean_co and d_name == clean_name and d_co == clean_co:
                is_duplicate = 1; break

        c_id = str(uuid.uuid4())
        c.execute("""
            INSERT INTO candidates (
                id, name_kr, email, phone, birth_year,
                current_title, sector, profile_summary,
                total_years, current_company,
                google_drive_url, raw_text, document_hash,
                is_duplicate, is_parsed, is_neo4j_synced, is_pinecone_synced,
                careers_json, education_json,
                has_big_company, has_startup, company_timeline,
                created_at, updated_at, source_file
            ) VALUES (
                ?,?,?,?,?,
                ?,?,?,
                ?,?,
                ?,?,?,
                ?,1,0,0,
                ?,?,
                ?,?,?,
                datetime('now'),datetime('now'),?
            )
        """, (
            c_id, name_kr, email, phone, birth_year,
            current_title, sector, summary,
            total_years, current_company,
            drive_link, text, doc_hash,
            is_duplicate,
            json.dumps(careers, ensure_ascii=False),
            json.dumps(edu, ensure_ascii=False),
            has_big, has_startup, company_timeline,
            filename,
        ))
        conn.commit()
        conn.close()

    if is_duplicate:
        return True, f"{name_kr} – 중복 저장"

    # Neo4j 동기화
    with neo_lock:
        with neo4j_driver.session() as session:
            session.run("""
                MERGE (c:Candidate {id: $id})
                SET c.name=$name, c.email=$email, c.phone=$phone,
                    c.current_company=$co, c.current_title=$title,
                    c.profile_summary=$summary,
                    c.total_years=$ty, c.sector=$sector
            """, id=c_id, name=name_kr, email=email, phone=phone,
                 co=current_company, title=current_title,
                 summary=summary, ty=total_years, sector=sector)

            for edge in parsed.get("neo4j_edges", []):
                act   = edge.get("action", "")
                skill = normalize_skill(edge.get("skill", ""))
                conf  = float(edge.get("confidence", 0.5))
                ev    = edge.get("evidence_span", "")
                if act and skill:
                    try:
                        session.run(f"""
                            MATCH (c:Candidate {{id: $id}})
                            MERGE (s:Skill {{name: $skill}})
                            MERGE (c)-[r:{act}]->(s)
                            SET r.confidence=$conf, r.evidence_span=$ev, r.source='v11'
                        """, id=c_id, skill=skill, conf=conf, ev=ev)
                    except Exception as e:
                        print(f"[Neo4j 스킬 오류] {skill}: {e}")

    # Pinecone 임베딩
    with pinecone_lock:
        try:
            emb_text = build_embedding_text({
                "name_kr": name_kr, "current_title": current_title,
                "sector": sector, "current_company": current_company,
                "profile_summary": summary, "raw_text": text,
            })
            chunks = chunk_text(emb_text)
            if chunks:
                emb_resp = openai_client.embeddings.create(
                    model="text-embedding-3-small", input=chunks
                )
                vectors = [
                    {
                        "id": f"{c_id}_chunk_{i}",
                        "values": e.embedding,
                        "metadata": {"candidate_id": c_id, "chunk_index": i},
                    }
                    for i, e in enumerate(emb_resp.data)
                ]
                pinecone_client.upsert(vectors, namespace="resume_vectors")
        except Exception as e:
            print(f"[Pinecone 오류] {filename}: {e}")

    return True, f"{name_kr} | {current_title} | {current_company} | {total_years}년 | {sector} | big={has_big} startup={has_startup}"

# ---------------------------------------------------------------------------
# 메인
# ---------------------------------------------------------------------------
def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--test",   action="store_true", help="20개만 테스트")
    parser.add_argument("--folder", type=str,            help="특정 폴더만")
    parser.add_argument("--workers", type=int, default=3, help="동시 처리 수")
    args = parser.parse_args()

    init_db()

    # 파일 수집 (변환본 우선, 중복 파일명 제거)
    dirs = [args.folder] if args.folder else TARGET_DIRS
    seen_names = set()
    files = []
    for d in dirs:
        if not os.path.isdir(d):
            continue
        for f in sorted(os.listdir(d)):
            if not f.lower().endswith((".pdf", ".docx", ".doc", ".txt")):
                continue
            if f.startswith("~$"):
                continue
            if os.path.getsize(os.path.join(d, f)) < 1024:
                continue
            if f not in seen_names:
                seen_names.add(f)
                files.append(os.path.join(d, f))

    print(f"대상 파일: {len(files)}개")
    if args.test:
        files = files[:20]
        print("테스트 모드: 20개")

    success = skipped = duplicated = 0

    with ThreadPoolExecutor(max_workers=args.workers) as executor:
        futures = {executor.submit(process_file, fp): os.path.basename(fp) for fp in files}
        for future in tqdm(as_completed(futures), total=len(futures)):
            fname = futures[future]
            try:
                ok, msg = future.result()
                if ok:
                    if "중복" in msg:
                        duplicated += 1
                        print(f"[DUP] {fname}: {msg}")
                    else:
                        success += 1
                        print(f"[OK ] {fname}: {msg}")
                else:
                    skipped += 1
                    if skipped <= 20 or "MD5" not in msg:
                        print(f"[SKP] {fname}: {msg}")
            except Exception as e:
                skipped += 1
                print(f"[ERR] {fname}: {e}")

    print(f"\n=== 완료 ===")
    print(f"신규: {success} | 중복: {duplicated} | 스킵/오류: {skipped}")

    with sqlite3.connect(DB_PATH) as conn:
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM candidates WHERE is_duplicate=0")
        print(f"DB 활성 후보자: {cur.fetchone()[0]}명")

if __name__ == "__main__":
    main()
