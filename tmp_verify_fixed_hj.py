import os
import json
import sys
from docx import Document

sys.path.append(os.getcwd())
from resume_parser import ResumeParser
from connectors.openai_api import OpenAIClient

def extract_text(filepath):
    doc = Document(filepath)
    text = []
    for p in doc.paragraphs:
        if p.text.strip(): text.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data: text.append(" | ".join(row_data))
    return "\n".join(text).strip()

if __name__ == "__main__":
    with open("secrets.json", "r") as f: secrets = json.load(f)
    client = OpenAIClient(secrets["OPENAI_API_KEY"])
    parser = ResumeParser(client)
    
    path = r"C:\Users\cazam\Downloads\02_resume 전처리\[NC] 현종민(IR)부문_86.docx"
    text = extract_text(path)
    print(f"--- EXTRACTED TEXT (LENGTH: {len(text)}) ---")
    print(text[:1000])
    
    print("\n--- NEW ANALYSIS ---")
    result = parser.parse(text)
    print(json.dumps(result, indent=2, ensure_ascii=False))
