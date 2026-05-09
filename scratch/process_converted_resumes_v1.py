import os
import sqlite3
import json
import hashlib
import re
import sys
import time
import io
import uuid
from docx import Document
import PyPDF2

# Add project root to path
PROJECT_ROOT = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템"
sys.path.append(PROJECT_ROOT)

from connectors.gemini_api import GeminiClient
from resume_parser import ResumeParser

# Configuration
RESUME_DIR = r"C:\Users\cazam\Downloads\02_resume_converted_v8"
DB_PATH = os.path.join(PROJECT_ROOT, "candidates.db")
SECRETS_PATH = os.path.join(PROJECT_ROOT, "secrets.json")

# Set encoding
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')

def calculate_hash(content_bytes):
    return hashlib.sha256(content_bytes).hexdigest()

def extract_text_docx(path):
    try:
        doc = Document(path)
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    try:
                        cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                        if cell_text:
                            row_text.append(cell_text)
                    except: pass
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    except Exception as e:
        print(f"  [DOCX Error] {os.path.basename(path)}: {e}", flush=True)
        return ""

def process_converted_files():
    print(f"[v1.4] Starting Process for Converted Resumes (.docx)...", flush=True)
    
    with open(SECRETS_PATH, "r", encoding="utf-8") as f:
        secrets = json.load(f)
        
    gemini = GeminiClient(secrets['GEMINI_API_KEY'])
    parser = ResumeParser(gemini)
    
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    
    # 1. Identify already processed files
    cur.execute("SELECT source_file FROM candidates WHERE source_file IS NOT NULL")
    existing_files = set(r[0] for r in cur.fetchall())
    
    # 2. Get existing names/hashes for duplication check
    cur.execute("SELECT name_kr, document_hash FROM candidates")
    existing_data = cur.fetchall()
    existing_names = set(r[0] for r in existing_data if r[0])
    existing_hashes = set(r[1] for r in existing_data if r[1])
    
    all_files = [f for f in os.listdir(RESUME_DIR) if f.lower().endswith('.docx')]
    
    new_candidates_files = []
    for f in all_files:
        if f in existing_files:
            continue
            
        # Match by name - more strict for converted files since many share substrings
        is_duplicate_name = False
        # Extract potential name from filename (e.g. "김철수_이력서.docx" -> "김철수")
        match = re.search(r'([가-힣]{2,4})', f)
        if match:
            extracted_name = match.group(1)
            if extracted_name in existing_names:
                is_duplicate_name = True
        
        if not is_duplicate_name:
            new_candidates_files.append(f)

    print(f"Initial filtering: {len(new_candidates_files)} potentially new candidates among {len(all_files)} files.", flush=True)
    
    processed = 0
    skipped = 0
    errors = 0
    
    # We will process in small batches or with a limit for now to show progress
    # User might want to run the whole thing. I'll do 20 for now.
    for i, filename in enumerate(new_candidates_files, 1):
        path = os.path.join(RESUME_DIR, filename)
        safe_filename = filename.encode('ascii', 'ignore').decode('ascii') or f"file_{i}"
        
        print(f"[{i}/{len(new_candidates_files)}] Processing: {safe_filename}", flush=True)
        
        try:
            with open(path, "rb") as f:
                content_bytes = f.read()
            f_hash = calculate_hash(content_bytes)
            
            if f_hash in existing_hashes:
                print(f"  [SKIP] Content hash already exists.", flush=True)
                cur.execute("UPDATE candidates SET source_file = ? WHERE document_hash = ?", (filename, f_hash))
                skipped += 1
                conn.commit()
                continue
            
            text = extract_text_docx(path)
            
            if not text or len(text) < 100:
                print(f"  [WARN] Text extraction weak ({len(text)} chars). Skipping.", flush=True)
                errors += 1
                continue
            
            parsed = parser.parse(text)
            if not parsed:
                print(f"  [FAIL] Gemini parsing failed.", flush=True)
                errors += 1
                continue
                
            if isinstance(parsed, list): parsed = parsed[0] if parsed else {}

            profile = parsed.get('candidate_profile', {})
            name_kr = profile.get('name_kr') or filename.split('_')[0].split('(')[0].strip()
            
            cid = str(uuid.uuid4())
            main_sectors = profile.get('main_sectors', [])
            sector_str = ", ".join(main_sectors) if isinstance(main_sectors, list) else str(main_sectors or '')
            exp_summary = profile.get('experience_summary', '')
            summary_str = "\n".join(exp_summary) if isinstance(exp_summary, list) else str(exp_summary or '')
            patterns = parsed.get('patterns', [])
            careers_json = json.dumps(patterns, ensure_ascii=False)
            
            cur.execute("""
                INSERT INTO candidates (
                    id, name_kr, raw_text, document_hash, is_parsed, 
                    sector, profile_summary, careers_json, total_years,
                    source_file, created_at, updated_at
                ) VALUES (?, ?, ?, ?, 1, ?, ?, ?, ?, ?, datetime('now'), datetime('now'))
            """, (
                cid, name_kr, text, f_hash, 
                sector_str, summary_str, careers_json, 
                profile.get('total_years_experience', 0),
                filename
            ))
            
            conn.commit()
            print(f"  [OK] Saved {name_kr}", flush=True)
            processed += 1
            existing_hashes.add(f_hash)
            existing_names.add(name_kr)
            
        except Exception as e:
            print(f"  [ERROR] {e}", flush=True)
            errors += 1
            
        time.sleep(1.0)
        
    conn.close()
    print(f"\nBatch Finished. Processed: {processed}, Skipped: {skipped}, Errors: {errors}", flush=True)

if __name__ == "__main__":
    process_converted_files()
