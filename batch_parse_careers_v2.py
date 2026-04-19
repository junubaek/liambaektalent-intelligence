import os
import json
import re
import glob
import time
import asyncio
from typing import Dict, List

from jd_compiler import client as genai_client
from google.genai import types
from neo4j import GraphDatabase

ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
CACHE_FILE = os.path.join(ROOT_DIR, "candidates_cache_jd.json")
SECRETS_FILE = os.path.join(ROOT_DIR, "secrets.json")

PDF_DIR = os.path.abspath(os.path.join(ROOT_DIR, "..", "02_resume 전처리"))
V8_DOCX_DIR = os.path.abspath(os.path.join(ROOT_DIR, "..", "02_resume_converted_v8"))

def load_secrets():
    with open(SECRETS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def save_cache(data: list):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def extract_text_from_file(clean_name: str) -> str:
    # 1. Search for .docx in v8 first
    import docx
    import pdfplumber
    
    docx_files = glob.glob(os.path.join(V8_DOCX_DIR, "**", f"*{clean_name}*.docx"), recursive=True)
    if not docx_files:
        # fuzzy search
        all_docx = glob.glob(os.path.join(V8_DOCX_DIR, "**", "*.docx"), recursive=True)
        docx_files = [f for f in all_docx if clean_name in os.path.basename(f)]
        
    if docx_files:
        try:
            doc = docx.Document(docx_files[0])
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Failed to read docx {docx_files[0]}: {e}")

    # 2. Search for PDF in 02_resume 전처리
    pdf_files = glob.glob(os.path.join(PDF_DIR, "**", f"*{clean_name}*.pdf"), recursive=True)
    if not pdf_files:
        all_pdfs = glob.glob(os.path.join(PDF_DIR, "**", "*.pdf"), recursive=True)
        pdf_files = [p for p in all_pdfs if clean_name in os.path.basename(p)]
        
    if pdf_files:
        try:
            with pdfplumber.open(pdf_files[0]) as pdf:
                # Up to 5 pages just to be safe
                text = "\n".join(page.extract_text() or '' for page in pdf.pages[:5])
                return text
        except Exception as e:
            print(f"Failed to read pdf {pdf_files[0]}: {e}")
            
    # 3. What if it's .doc? We should try to run the converter
    doc_files = glob.glob(os.path.join(PDF_DIR, "**", f"*{clean_name}*.doc"), recursive=True)
    if not doc_files:
        all_docs = glob.glob(os.path.join(PDF_DIR, "**", "*.doc"), recursive=True)
        doc_files = [p for p in all_docs if clean_name in os.path.basename(p) and not os.path.basename(p).startswith("~$")]
        
    if doc_files:
        print(f"  [Auto-Convert] Found .doc for {clean_name}. Converting to .docx...")
        try:
            from scripts.batch_convert_doc_v8 import convert_doc_to_docx
            # batch_convert_doc_v8 converts EVERYTHING in PDF_DIR, which might be overkill, 
            # but it is idempotent (skips converted). So let's run it.
            convert_doc_to_docx()
            
            # recursive check now
            docx_files = glob.glob(os.path.join(V8_DOCX_DIR, "**", f"*{clean_name}*.docx"), recursive=True)
            if not docx_files:
                all_docx = glob.glob(os.path.join(V8_DOCX_DIR, "**", "*.docx"), recursive=True)
                docx_files = [f for f in all_docx if clean_name in os.path.basename(f)]
            
            if docx_files:
                doc = docx.Document(docx_files[0])
                text = "\n".join([para.text for para in doc.paragraphs])
                return text
        except Exception as e:
            print(f"Failed to auto-convert doc {doc_files[0]}: {e}")

    return ""

def smart_chunk(raw_text: str) -> str:
    # Remove self-intro parts and focus on Career
    clean_text = ' '.join(raw_text.split())
    
    # Locate index of Career words
    keywords = ["경력사항", "경력", "업무경험", "Experience", "Work Experience", "프로젝트 이력"]
    start_idx = -1
    for kw in keywords:
        idx = clean_text.find(kw)
        if idx != -1:
            if start_idx == -1 or idx < start_idx:
                start_idx = idx
                
    if start_idx != -1:
        # Give a small 200 character buffer before the keyword just in case
        actual_start = max(0, start_idx - 200)
        chunk = clean_text[actual_start:]
    else:
        # If no keywords found, maybe the whole document is short
        chunk = clean_text

    # Skip specific sections by regex
    chunk = re.sub(r'자기소개서.*?(경력|$)', r'\1', chunk, flags=re.IGNORECASE | re.DOTALL)
    chunk = re.sub(r'성장과정.*?(경력|$)', r'\1', chunk, flags=re.IGNORECASE | re.DOTALL)
    chunk = re.sub(r'지원동기.*?(경력|$)', r'\1', chunk, flags=re.IGNORECASE | re.DOTALL)

    # Return up to 6000 chars for extreme depth
    return chunk[:6000]

async def process_candidate(session, cand: dict, driver) -> bool:
    cand_id = cand.get("id")
    raw_name = cand.get("name", "")
    phone = cand.get("phone", "")
    clean_name = cand.get("name_kr", "") or re.sub(r'\[.*?\]|\(.*?\)|\..*', '', raw_name).replace('부문', '').replace('원본', '').strip()
    
    # 1. Deduplciation checking (done outside, but double check)
    if cand.get("_duplicate_skipped"):
        return False

    # 2. Extract Document Text
    raw_text = extract_text_from_file(clean_name)
    if not raw_text or len(raw_text) < 50:
        raw_text = cand.get("summary", "")
        # fallback to the 500 length summary
        if not raw_text:
            return False

    chunked_text = smart_chunk(raw_text)
    
    system_prompt = """당신은 최고 수준의 데이터 정밀도를 지향하는 리크루팅 어시스턴트입니다. 
아래의 비정형 이력 텍스트를 분석하여 구조화된 JSON 상태로 변환하세요.

[무결성 가이드라인]
1. 연락처/기본정보 추출: 이력서 내에 이메일, 전화번호(010-), 연차(seniority)가 명시되어 있다면 해당 필드에 정확히 담아주세요. 없다면 빈 문자열("")로 반환하세요. 연차는 'Junior', 'Middle', 'Senior' 중 하나로 분류하세요.
2. 날짜 무결성 (가장 중요): 근무 기간(연/월)이 기재된 경우 '단 하나도 빠짐없이 전부' 추출하여 period 필드에 기재하세요. (예: '2022.08 ~ 현재', '2015.07 ~ 2022.07').
3. 직급 표준화: 원문의 직급/직무를 기재하세요.
4. 전체 경력 추출 (가장 중요): 제공된 이력서 텍스트 안에 회사가 10개라면 10개 전부를 빠짐없이 careers 리스트에 담아 추출하세요. 자기소개서 내용이 섞여있다면 직무 관련 경력만 뽑고 무시하세요.

[parsed_career_json 규칙]
- 동일한 company명은 반드시 1개 객체로만 기록. 여러 프로젝트가 있어도 동일 회사면 하나로 병합.
- 배열 내 company 중복 절대 금지.
- 최대 항목 수: 실제 재직 회사 수와 동일하게.

반드시 다음 JSON 스키마를 준수해야 합니다.
{
    "email": "string",
    "phone": "string",
    "seniority": "string",
    "careers": [
        {"company": "string", "team": "string", "position": "string", "period": "string"}
    ]
}"""

    config = types.GenerateContentConfig(
        response_mime_type="application/json",
        system_instruction=system_prompt,
        temperature=0.0
    )
    
    try:
        response = await genai_client.aio.models.generate_content(
            model="gemini-2.5-flash",
            contents=[f'분석 대상 텍스트: "{chunked_text}"'],
            config=config
        )
        parsed_data = json.loads(response.text)
        careers = parsed_data.get("careers", [])
        
        # Inject parsed data to cand dict
        cand["parsed_career_json"] = careers
        if parsed_data.get("email"): cand["email"] = parsed_data["email"]
        if parsed_data.get("phone"): cand["phone"] = parsed_data["phone"]
        if parsed_data.get("seniority"): cand["seniority"] = parsed_data["seniority"]
        
        # Save to neo4j
        career_str = json.dumps(careers, ensure_ascii=False)
        query = "MATCH (c:Candidate {id: $id}) SET c.parsed_career_json = $data"
        
        # Add optional metadata sets
        update_clauses = []
        params = {"id": cand_id, "data": career_str}
        
        if parsed_data.get("email"): 
            update_clauses.append("c.email = $email")
            params["email"] = parsed_data["email"]
        if parsed_data.get("phone"): 
            update_clauses.append("c.phone = $phone")
            params["phone"] = parsed_data["phone"]
            
        if update_clauses:
            query += ", " + ", ".join(update_clauses)
            
        session.run(query, **params)
        
        print(f"✅ Success: {clean_name} ({len(careers)} companies extracted)")
        return True
    except Exception as e:
        print(f"❌ Gemini Error for {clean_name}: {e}")
        return False

async def main():
    print("🚀 Starting Batch Parser V2...")
    secrets = load_secrets()
    neo4j_pwd = secrets.get('NEO4J_PASSWORD', 'toss1234')
    driver = GraphDatabase.driver('bolt://localhost:7687', auth=('neo4j', neo4j_pwd))
    
    data = load_cache()
    
    # De-Duplication Step
    print("🔍 Scanning for duplicates by name + phone/email...")
    seen = {}
    duplicates = 0
    for cand in data:
        # Exclude those with no clean info
        raw_name = cand.get("name", "")
        clean_name = cand.get("name_kr", "") or re.sub(r'\[.*?\]|\(.*?\)|\..*', '', raw_name).replace('부문', '').replace('원본', '').strip()
        phone = cand.get("phone", "").replace("-", "").strip()
        email = cand.get("email", "").strip()
        
        if not clean_name: continue
        
        # Identifier: strict name + (phone OR email). Only identical if contact info matches.
        # If both phone and email are blank, we skip deduplication to avoid nuking unrelated namesakes.
        if phone or email:
            key = f"{clean_name}|{phone}|{email}"
            if key in seen:
                cand["_duplicate_skipped"] = True
                duplicates += 1
            else:
                seen[key] = True
                
    print(f"   -> Found {duplicates} duplicated rows in cache. They will be skipped.")
    
    targets = [c for c in data if not c.get("_duplicate_skipped") and "parsed_career_json" not in c]
    print(f"📝 Total unparsed valid candidates: {len(targets)}")
    
    with driver.session() as session:
        for i, cand in enumerate(targets):
            raw_name = cand.get("name", "")
            print(f"[{i+1}/{len(targets)}] Processing {raw_name}...")
            
            success = await process_candidate(session, cand, driver)
            
            # Periodically write back to cache to prevent data loss on crash
            if i > 0 and i % 5 == 0:
                save_cache(data)
                
            # Rate Limiting
            await asyncio.sleep(1.0)
            
    # Final save
    save_cache(data)
    driver.close()
    print("🏁 Batch Processing Complete!")

if __name__ == "__main__":
    asyncio.run(main())
