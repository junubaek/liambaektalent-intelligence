import os
import json
import time
import uuid
import sqlite3
import hashlib
import re
import sys
from datetime import datetime
import pdfplumber
from docx import Document
import google.generativeai as genai
from neo4j import GraphDatabase

from connectors.gdrive_api import GDriveConnector
from googleapiclient.http import MediaFileUpload
from openai import OpenAI
from connectors.pinecone_api import PineconeClient
from batch_pinecone_sync import chunk_text

sys.stdout.reconfigure(encoding='utf-8')

DB_PATH = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\candidates.db"
secrets_path = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\secrets.json"

with open(secrets_path, "r", encoding="utf-8") as f:
    secrets = json.load(f)

genai.configure(api_key=secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel("gemini-2.5-flash")

n_uri = secrets.get("NEO4J_URI", "neo4j://127.0.0.1:7687")
n_user = secrets.get("NEO4J_USERNAME", "neo4j")
n_pw = secrets.get("NEO4J_PASSWORD", "toss1234")
driver = GraphDatabase.driver(n_uri, auth=(n_user, n_pw))

gdrive = GDriveConnector()
folder_id = secrets.get("GOOGLE_DRIVE_FOLDER_ID")

openai_client = OpenAI(api_key=secrets.get("OPENAI_API_KEY", ""))
pc_host = secrets.get("PINECONE_HOST", "").rstrip("/")
if not pc_host.startswith("https://"): pc_host = f"https://{pc_host}"
pinecone_client = PineconeClient(secrets.get("PINECONE_API_KEY", ""), pc_host)

# 13 specific files
TARGET_FILES = [
    "[리벨리온] 황승현(NPU Runtime Software Engineer)부문.pdf",
    "[리벨리온] 홍용기(NPU Library Software Engineer)부문.pdf",
    "[리벨리온] 최현석(NPU Runtime Software Engineer)부문.pdf",
    "[리벨리온] 최성우(NPU Runtime Software Engineer)부문.pdf",
    "[리벨리온] 정태영(NPU Runtime Software Engineer)부문.pdf",
    "[리벨리온] 전형준(Moreh Software Engineer)부문.pdf",
    "[리벨리온] 전예찬(NPU Library Software Engineer)부문.pdf",
    "[리벨리온] 이석현(NPU Runtime Software Engineer)부문.pdf",
    "[리벨리온] 송우석(NPU Library Software Engineer)부문.pdf",
    "[리벨리온] 박지민(NPU Runtime Software Engineer (LLM Serving))부문.pdf",
    "[리벨리온] 김학주(NPU Runtime Software Engineer)부문.docx",
    "[리벨리온] 김정근(NPU Library Software Engineer)부문.docx",
    "[리벨리온] 김기덕(NPU Runtime Software Engineer)부문.pdf"
]

TARGET_DIR = r"C:\Users\cazam\Downloads\02_resume 전처리"

MEGA_PROMPT = """이력서 텍스트를 분석해서 아래 구조를 JSON Object 형식으로 출력해. 반드시 코드블럭 없이 JSON만 반환할 것.

{{
  "name_kr": "본문에서 순수 본명만 추출 (직무 꼬리표 붙이지 말 것). 못 찾으면 null",
  "phone": "010-XXXX-XXXX 형태 숫자만. 못 찾으면 null",
  "email": "이메일주소. 못 찾으면 null",
  "birth_year": 출생연도 숫자 4자리 (예: 1990). 못 찾으면 null,
  "summary": "주요 경력 2줄 요약 (전화번호/이메일/주소/생년월일 등 개인정보 절대 포함 금지!!!)",
  "sector": "이 후보자는 NPU / Library / Runtime / Software Engineer이므로 반드시 ['Eng_AI', 'Eng_Semi', 'Eng_Embedded'] 중 1개 혹은 2개를 조합(쉼표 구분)하여 반환해야 함. 예: 'Eng_AI' 또는 'Eng_Semi, Eng_AI' 또는 'Eng_Embedded'",
  "education_json": [
    {{"school": "대학교 이상 학교명", "major": "전공", "degree": "학사/석사 등", "year": "졸업년도"}}
  ],
  "careers_json": [
    {{
      "company": "회사명",
      "title": "직책/부서",
      "start_date": "YYYY.MM",
      "end_date": "YYYY.MM (또는 현재)"
    }}
  ],
  "neo4j_edges": [
    {{
      "action": "BUILT|DESIGNED|MANAGED|ANALYZED|LAUNCHED|NEGOTIATED|GREW|SUPPORTED 중 택1",
      "skill": "아래 지시된 Skill 목록 중 택1",
      "confidence": 0.0 ~ 1.0,
      "evidence_span": "해당 판단을 내리게 된 본문 문구 복사"
    }}
  ]
}}

[Skill 목록]:
Payment_and_Settlement_System, Service_Planning, Product_Manager, Data_Pipeline_Construction, Data_Engineering, Data_Analysis,
Backend, Frontend, Machine_Learning, MLOps, DevOps, Financial_Accounting, Corporate_Strategic_Planning, 사업개발_BD, 퍼포먼스마케팅,
Treasury_Management, FX_Dealing, Corporate_Funding, IPO_Preparation_and_Execution, Recruiting_and_Talent_Acquisition,
Organizational_Development, B2B영업, 물류_Logistics, Backend_Python, Backend_Java, Backend_Go, Backend_Node,
Kubernetes, Infrastructure_and_Cloud, 보안_Security, FinTech, Natural_Language_Processing, 컴퓨터비전_CV, 추천시스템,
Deep_Learning, Corporate_Legal_Counsel, Intellectual_Property, Legal_Compliance, Contract_Management, Litigation

이력서 본문:
{text}
"""

def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    text = ""
    try:
        import fitz
        if ext == "pdf":
            with fitz.open(filepath) as doc:
                for page in doc: text += page.get_text()
        elif ext in ("docx", "doc"):
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            for t in doc.tables:
                for r in t.rows:
                    for ce in r.cells:
                        if ce.text: text += ce.text + " "
    except Exception as e:
        print(f"Extraction error for {filepath}: {e}")
    return text.strip()

def calculate_career_stats(careers):
    if not careers: return "미상", 0.0
    total_months = 0
    latest_end_dt = datetime(1900, 1, 1)
    curr_comp = "미상"
    
    def parse_dt(d_str):
        if not d_str or '현재' in d_str or '재직' in d_str or 'ing' in d_str.lower(): return datetime.now()
        m = re.findall(r'(\d{4})[^\d]*(\d{1,2})', d_str)
        if m: return datetime(int(m[0][0]), int(m[0][1]), 1)
        m2 = re.findall(r'(\d{4})', d_str)
        if m2: return datetime(int(m2[0]), 1, 1)
        return datetime.now()

    for c in careers:
        st = parse_dt(c.get('start_date', ''))
        ed = parse_dt(c.get('end_date', ''))
        if ed < st: ed = st
        months = (ed.year - st.year) * 12 + ed.month - st.month
        total_months += max(months, 0)
        
        if ed >= latest_end_dt:
            latest_end_dt = ed
            if c.get('company'): curr_comp = c.get('company')
            
    total_y = round(total_months / 12.0, 1)
    return curr_comp, total_y

def extract_fallback_name(filename):
    m = re.search(r'[가-힣]{2,4}', filename.replace("이력서", "").replace("포트폴리오", "").replace("개발자", ""))
    return m.group(0) if m else "미상"

def clean_existing_candidate(name, conn):
    c = conn.cursor()
    # Find existing IDs
    c.execute("SELECT id FROM candidates WHERE name_kr LIKE ?", (f"%{name}%",))
    rows = c.fetchall()
    ids = [r[0] for r in rows]
    
    if ids:
        print(f"[{name}] Cleaning up existing candidate records (IDs: {ids})")
        for cid in ids:
            # Delete from SQLite
            c.execute("DELETE FROM candidates WHERE id = ?", (cid,))
            
            # Delete from Neo4j
            try:
                with driver.session() as session:
                    session.run("MATCH (c:Candidate {id: $id}) DETACH DELETE c", id=cid)
            except Exception as e:
                print(f"Neo4j delete error for {cid}: {e}")
                
            # Delete from Pinecone
            try:
                chunk_ids = [f"{cid}_chunk_{i}" for i in range(100)]
                pinecone_client.delete(ids=chunk_ids, namespace="resume_vectors")
            except Exception as e:
                print(f"Pinecone delete error for {cid}: {e}")
        conn.commit()

def force_process_file(filename):
    filepath = os.path.join(TARGET_DIR, filename)
    if not os.path.exists(filepath):
        print(f"[ERROR] File does not exist: {filepath}")
        return False
        
    print(f"\n>>> Processing: {filename}")
    text = extract_text(filepath)
    if len(text) < 50:
        print(f"[ERROR] Text too short: {len(text)}")
        return False
        
    doc_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    fallback_name = extract_fallback_name(filename)
    
    # Clean up existing candidate records first to prevent duplicates!
    conn = sqlite3.connect(DB_PATH)
    clean_existing_candidate(fallback_name, conn)
    
    # Upload to Google Drive (reuse or upload)
    drive_link = None
    try:
        drive_res = gdrive.service.files().list(q=f"name='{filename}' and '{folder_id}' in parents and trashed=false", fields="files(id, webViewLink)").execute()
        if drive_res.get('files'):
            drive_link = drive_res['files'][0]['webViewLink']
            print(f"  [Drive] Found existing file: {drive_link}")
        else:
            ext = filename.lower().split('.')[-1]
            mt = 'application/pdf' if ext == 'pdf' else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            media = MediaFileUpload(filepath, mimetype=mt, resumable=True)
            uploaded = gdrive.service.files().create(body={'name': filename, 'parents': [folder_id]}, media_body=media, fields='webViewLink').execute()
            drive_link = uploaded.get('webViewLink')
            print(f"  [Drive] Uploaded new file: {drive_link}")
    except Exception as e:
        print(f"  [Drive] Upload Error: {e}")

    # Gemini Parsing
    parsed = None
    for attempt in range(3):
        try:
            prompt = MEGA_PROMPT.replace("{text}", f"[파일명: {filename}]\n\n" + text[:6000])
            res = model.generate_content(prompt)
            raw = res.text.replace("```json", "").replace("```", "").strip()
            parsed = json.loads(raw)
            break
        except Exception as e:
            print(f"  [Gemini] Attempt {attempt+1} failed: {e}")
            time.sleep(2)
            
    if not parsed:
        print("[ERROR] Gemini Parsing Failed")
        conn.close()
        return False

    name_kr = parsed.get("name_kr")
    if not name_kr or name_kr == "null": name_kr = fallback_name

    email = parsed.get("email", "")
    phone = parsed.get("phone", "")
    birth_year = parsed.get("birth_year", 0)
    sector = parsed.get("sector", "Eng_AI")
    summary = parsed.get("summary", "")
    careers = parsed.get("careers_json", [])
    edu = parsed.get("education_json", [])
    
    # Strictly verify sector is in AI, Semi, Embedded
    allowed_sectors = ['Eng_AI', 'Eng_Semi', 'Eng_Embedded']
    sectors_split = [s.strip() for s in sector.split(',')]
    verified_sectors = [s for s in sectors_split if s in allowed_sectors]
    if not verified_sectors:
        # Fallback default
        if "library" in filename.lower() or "runtime" in filename.lower() or "moreh" in filename.lower():
            sector = "Eng_AI, Eng_Semi"
        else:
            sector = "Eng_AI"
    else:
        sector = ", ".join(verified_sectors)
        
    current_company, total_years = calculate_career_stats(careers)
    c_id = str(uuid.uuid4())
    
    # Save to SQLite
    c = conn.cursor()
    c.execute('''
        INSERT INTO candidates (
            id, name_kr, email, phone, birth_year, sector, profile_summary, total_years, current_company,
            google_drive_url, raw_text, document_hash, is_duplicate, is_parsed, is_neo4j_synced, is_pinecone_synced,
            careers_json, education_json, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, 1, 1, 1, ?, ?, datetime('now'), datetime('now'))
    ''', (c_id, name_kr, email, phone, birth_year, sector, summary, total_years, current_company,
          drive_link, text, doc_hash, json.dumps(careers, ensure_ascii=False), json.dumps(edu, ensure_ascii=False)))
    conn.commit()
    conn.close()
    
    print(f"  [SQLite] Saved: {name_kr} | Sector: {sector} | Company: {current_company}")

    # Neo4j Sync
    try:
        with driver.session() as session:
            session.run("""
                MERGE (c:Candidate {id: $id})
                SET c.name = $name_kr, c.phone = $phone, c.email = $email,
                    c.current_company = $current_company, c.profile_summary = $summary,
                    c.total_years = $total_years, c.sector = $sector
            """, id=c_id, name_kr=name_kr, phone=phone, email=email, 
                 current_company=current_company, summary=summary, total_years=total_years, sector=sector)
                 
            for edge in parsed.get("neo4j_edges", []):
                act, skill, conf, ev = edge.get("action", ""), edge.get("skill", ""), float(edge.get("confidence", 0.5)), edge.get("evidence_span", "")
                if act and skill:
                    session.run(f"""
                        MERGE (c:Candidate {{id: $id}})
                        MERGE (s:Skill {{name: $skill}})
                        MERGE (c)-[r:{act}]->(s)
                        SET r.confidence = $conf, r.evidence_span = $ev, r.source = 'v10_ingest'
                    """, id=c_id, skill=skill, conf=conf, ev=ev)
        print("  [Neo4j] Synced")
    except Exception as e:
        print(f"  [Neo4j] Sync Error: {e}")

    # Pinecone Sync
    try:
        chunks = chunk_text(text)
        if chunks:
            response = openai_client.embeddings.create(model="text-embedding-3-small", input=chunks)
            vectors = []
            for i, emb in enumerate(response.data):
                vectors.append({
                    "id": f"{c_id}_chunk_{i}",
                    "values": emb.embedding,
                    "metadata": {"candidate_id": c_id, "chunk_index": i}
                })
            pinecone_client.upsert(vectors, namespace="resume_vectors")
            print("  [Pinecone] Upserted chunks")
    except Exception as e:
        print(f"  [Pinecone] Sync Error: {e}")
        
    return True

def main():
    print(f"Forced ingestion of {len(TARGET_FILES)} NPU/Library candidates...")
    success_count = 0
    for filename in TARGET_FILES:
        ok = force_process_file(filename)
        if ok:
            success_count += 1
    print(f"\nFinished: {success_count}/{len(TARGET_FILES)} files processed successfully.")

if __name__ == "__main__":
    main()
