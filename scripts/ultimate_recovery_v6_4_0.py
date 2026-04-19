import os
import json
import time
import requests
import sqlite3
from docx import Document
import PyPDF2

# Import project modules
import sys
sys.path.append(os.getcwd())
from connectors.openai_api import OpenAIClient
from resume_parser import ResumeParser
from headhunting_engine.normalization.pattern_merger import PatternMerger
from headhunting_engine.matching.scorer import Scorer
from connectors.gdrive_api import GDriveConnector

# Configuration
DIR_RAW = r"C:\Users\cazam\Downloads\02_resume 전처리"
DIR_CONV = r"C:\Users\cazam\Downloads\02_resume_converted_docx"
GDRIVE_FOLDER_ID = "1VzJEeoXG239PVR3IoJM5jC28KccMdkth"

with open("secrets.json", "r") as f:
    secrets = json.load(f)

NOTION_TOKEN = secrets["NOTION_API_KEY"]
NOTION_DB_ID = secrets["NOTION_DATABASE_ID"]

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    text = []
    try:
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
        elif ext == '.docx':
            doc = Document(filepath)
            for p in doc.paragraphs:
                if p.text.strip(): text.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data: text.append(" | ".join(row_data))
    except Exception as e:
        print(f"  ⚠️ Error extracting {filepath}: {e}")
    return "\n".join(text).strip()

class IntelligentRecovery:
    def __init__(self):
        self.headers = {
            "Authorization": f"Bearer {NOTION_TOKEN}",
            "Content-Type": "application/json",
            "Notion-Version": "2022-06-28"
        }
        self.openai_client = OpenAIClient(secrets["OPENAI_API_KEY"])
        self.parser = ResumeParser(self.openai_client)
        self.merger = PatternMerger()
        self.scorer = Scorer()
        self.gdrive = GDriveConnector()

    def get_clean_notion_map(self):
        print("📋 Fetching LIVE Notion Hub map (Archived records excluded)...")
        url = f"https://api.notion.com/v1/databases/{NOTION_DB_ID}/query"
        pages = []
        has_more = True
        start_cursor = None
        
        while has_more:
            payload = {"page_size": 100}
            if start_cursor: payload["start_cursor"] = start_cursor
            resp = requests.post(url, headers=self.headers, json=payload).json()
            results = resp.get('results', [])
            for r in results:
                if not r.get('archived'):
                    pages.append(r)
            
            has_more = resp.get('has_more', False)
            start_cursor = resp.get('next_cursor')
            print(f"  -> Loaded {len(pages)} live pages...")
            
        n_map = {}
        for p in pages:
            props = p.get('properties', {})
            title = props.get('이름', {}).get('title', [])
            if title:
                name = title[0].get('plain_text')
                # Check for blanks in specific fields
                has_patterns = len(props.get('Experience Patterns', {}).get('multi_select', [])) > 0
                
                if name not in n_map:
                    n_map[name] = []
                n_map[name].append({
                    "id": p['id'],
                    "has_data": has_patterns
                })
        return n_map

    def process_candidate(self, name_base, filepath, target_id):
        try:
            print(f"🔄 Processing: {name_base}")
            text = extract_text(filepath)
            
            if len(text) < 300:
                print(f"  ⚠️ Skipping {name_base}: Text too short ({len(text)} chars).")
                return

            # AI Intelligence
            intel = self.parser.parse(text)
            
            # Pattern Merging
            raw_patterns = [p.get("pattern", "") for p in intel.get("patterns", [])]
            consolidated = self.merger.merge_list(raw_patterns, limit=7)
            notion_patterns = [{"name": p[:100]} for p in consolidated]
            
            # Scoring
            score, _ = self.scorer.calculate_score(intel.get("patterns", []), set(), set(), intel)
            
            # GDrive Link
            _, drive_link = self.gdrive.upload_file(filepath, GDRIVE_FOLDER_ID)
            
            # Properties
            profile = intel.get("candidate_profile", {})
            summary = profile.get("experience_summary", "") or intel.get("summary", "")
            
            sectors = profile.get("primary_sector", "Unclassified")
            if isinstance(sectors, str): sectors = [sectors]
            multi_sectors = [{"name": s} for s in sectors]

            props = {
                "Primary Sector": {"multi_select": multi_sectors},
                "Experience Patterns": {"multi_select": notion_patterns},
                "Trajectory Grade": {"select": {"name": intel.get("career_path_quality", {}).get("trajectory_grade", "Neutral")}},
                "v6.2 Score": {"number": score},
                "구글드라이브 링크": {"url": drive_link}
            }
            if summary:
                props["경력 Summary"] = {"rich_text": [{"text": {"content": summary[:2000]}}]}
            
            requests.patch(f"https://api.notion.com/v1/pages/{target_id}", headers=self.headers, json={"properties": props}, timeout=15)
            print(f"  ✅ Successfully Optimized: {name_base}")

        except Exception as e:
            print(f"  ❌ Error in {name_base}: {e}")

    def run(self):
        notion_map = self.get_clean_notion_map()
        
        all_files = {}
        for d in [DIR_RAW, DIR_CONV]:
            if not os.path.exists(d): continue
            for f in os.listdir(d):
                if f.lower().endswith(('.pdf', '.docx')) and not f.startswith("~$"):
                    name = os.path.splitext(f)[0]
                    if name not in all_files:
                        all_files[name] = os.path.join(d, f)
        
        print(f"📂 Found {len(all_files)} files in repository.")
        
        to_process = []
        for name, files in notion_map.items():
            for entry in files:
                if not entry['has_data']:
                    if name in all_files:
                        to_process.append((name, all_files[name], entry['id']))
        
        print(f"🚀 Found {len(to_process)} blank live entries to fill.")
        
        for i, (name, path, nid) in enumerate(to_process):
            self.process_candidate(name, path, nid)
            if (i+1) % 5 == 0:
                print(f"🏁 Checkpoint: {i+1}/{len(to_process)} optimized.")
            time.sleep(1)

if __name__ == "__main__":
    recovery = IntelligentRecovery()
    recovery.run()
