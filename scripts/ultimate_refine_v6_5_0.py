import os
import json
import time
import requests
import re
from docx import Document
import PyPDF2

# Project Imports
import sys
sys.path.append(os.getcwd())
from connectors.openai_api import OpenAIClient
from resume_parser import ResumeParser
from headhunting_engine.normalization.pattern_merger import PatternMerger
from headhunting_engine.matching.scorer import Scorer
from connectors.gdrive_api import GDriveConnector

# Configuration
DIR_RAW = r"C:\Users\cazam\Downloads\02_resume 전처리"
DIR_CONV = r"C:\Users\cazam\Downloads\02_resume_converted_docx"
GDRIVE_FOLDER_ID = "1VzJEeoXG239PVR3IoJM5jC28KccMdkth"

with open("secrets.json", "r") as f:
    secrets = json.load(f)

NOTION_TOKEN = secrets["NOTION_API_KEY"]
NOTION_DB_ID = secrets["NOTION_DATABASE_ID"]

headers = {
    "Authorization": f"Bearer {NOTION_TOKEN}",
    "Content-Type": "application/json",
    "Notion-Version": "2022-06-28"
}

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    text = []
    try:
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
        elif ext == '.docx':
            doc = Document(filepath)
            # 1. Paragraphs
            for p in doc.paragraphs:
                if p.text.strip(): text.append(p.text.strip())
            # 2. Tables (Crucial for structured resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_content:
                        text.append(" | ".join(dict.fromkeys(row_content))) # Dedupe within row
    except Exception as e:
        print(f"  ⚠️ Error extracting {filepath}: {e}")
    return "\n".join(text).strip()

def jaccard_similarity(set1, set2):
    if not set1 and not set2: return 1.0
    if not set1 or not set2: return 0.0
    intersection = len(set1.intersection(set2))
    union = len(set1.union(set2))
    return intersection / union

def get_core_name(full_name):
    # 1. Remove [Company] or (Role) blocks first
    name = re.sub(r'\[.*?\]', '', full_name)
    name = re.sub(r'\(.*?\)', '', name)
    
    # 2. Remove common suffixes and noisy characters
    name = re.sub(r'\s|부문|_원본|_이력서|원본|_포트폴리오|포트폴리오|이력서|지원서', '', name)
    
    # 3. Handle specific patterns like "JY_" or "Android Developer_"
    name = re.sub(r'^[A-Z]{2,}_', '', name)
    name = re.sub(r'^[A-Za-z\s]+_', '', name)
    
    # 4. Remove year of birth (e.g. 87, 1987)
    name = re.sub(r'\d{2,4}', '', name)
    
    # 5. Extract the first block of Korean characters (usually 2-4 chars)
    match = re.search(r'[가-힣]{2,4}', name)
    if match:
        return match.group(0)
    
    return name.strip()

class UltimateRefiner:
    def __init__(self):
        self.openai_client = OpenAIClient(secrets["OPENAI_API_KEY"])
        self.parser = ResumeParser(self.openai_client)
        self.merger = PatternMerger()
        self.scorer = Scorer()
        self.gdrive = GDriveConnector()

    def get_live_pages(self):
        print("📋 Fetching all live pages from Notion...")
        url = f"https://api.notion.com/v1/databases/{NOTION_DB_ID}/query"
        pages = []
        has_more = True
        start_cursor = None
        while has_more:
            payload = {"page_size": 100}
            if start_cursor: payload["start_cursor"] = start_cursor
            resp = requests.post(url, headers=headers, json=payload).json()
            results = resp.get('results', [])
            for r in results:
                if not r.get('archived'):
                    pages.append(r)
            has_more = resp.get('has_more', False)
            start_cursor = resp.get('next_cursor')
            print(f"  -> Loaded {len(pages)} live pages...")
        return pages

    def run_aggressive_dedupe(self, pages):
        print("\n🧐 Phase 1: Aggressive Deduplication (Threshold: 0.3/0.2)...")
        name_groups = {}
        for p in pages:
            props = p['properties']
            name_list = props.get('이름', {}).get('title', [])
            if not name_list: continue
            name = name_list[0].get('plain_text')
            core = get_core_name(name)
            
            if core not in name_groups: name_groups[core] = []
            
            patterns = set([pt["name"] for pt in props.get('Experience Patterns', {}).get('multi_select', [])])
            summary = "".join([t["plain_text"] for t in props.get("경력 Summary", {}).get("rich_text", [])])
            sector = [s["name"] for s in props.get('Primary Sector', {}).get('multi_select', [])]
            
            name_groups[core].append({
                "id": p['id'],
                "name": name,
                "patterns": patterns,
                "summary": summary,
                "sector": sector,
                "score": props.get('v6.2 Score', {}).get('number', 0) or 0,
                "has_link": 1 if props.get('구글드라이브 링크', {}).get('url') else 0
            })

        archived = 0
        for core, group in name_groups.items():
            if len(group) <= 1: continue
            group.sort(key=lambda x: (x['has_link'], x['score'], len(x['patterns']), len(x['summary'])), reverse=True)
            to_keep = [group[0]]
            for i in range(1, len(group)):
                curr = group[i]
                is_dupe = False
                for kept in to_keep:
                    sim_p = jaccard_similarity(curr['patterns'], kept['patterns'])
                    common_k = set(curr['summary'].split()).intersection(set(kept['summary'].split()))
                    sim_s = len(common_k) / max(len(curr['summary'].split()), 1)
                    
                    # Aggressive thresholds
                    if curr['sector'] == kept['sector'] and (sim_p > 0.3 or sim_s > 0.2):
                        is_dupe = True
                        break
                if is_dupe:
                    print(f"  🗑️ Archiving Dupe: {curr['name']} ({curr['id']})")
                    requests.patch(f"https://api.notion.com/v1/pages/{curr['id']}", headers=headers, json={"archived": True})
                    archived += 1
                else:
                    to_keep.append(curr)
        print(f"✅ Aggressive Dedupe Complete: {archived} archived.")

    def run_fuzzy_recovery(self):
        print("\n📂 Phase 2: Fuzzy Name Matching & Short Text Recovery (Min: 100 chars)...")
        # Re-fetch live pages after dedupe
        pages = self.get_live_pages()
        
        # Build file index
        file_repo = {}
        for d in [DIR_RAW, DIR_CONV]:
            if not os.path.exists(d): continue
            for f in os.listdir(d):
                if f.lower().endswith(('.pdf', '.docx')) and not f.startswith("~$"):
                    fname = os.path.splitext(f)[0]
                    core = get_core_name(fname)
                    if core not in file_repo: file_repo[core] = []
                    file_repo[core].append(os.path.join(d, f))

        to_fill = []
        for p in pages:
            props = p['properties']
            has_data = len(props.get('Experience Patterns', {}).get('multi_select', [])) > 0
            if not has_data:
                name_list = props.get('이름', {}).get('title', [])
                full_name = name_list[0]['plain_text'] if name_list else "Unknown"
                core = get_core_name(full_name)
                if core in file_repo:
                    to_fill.append((full_name, file_repo[core][0], p['id']))

        print(f"🚀 Found {len(to_fill)} blanks with potential fuzzy matches.")
        
        for i, (name, path, nid) in enumerate(to_fill):
            try:
                print(f"🔄 [{i+1}/{len(to_fill)}] Recovering: {name} via {os.path.basename(path)}")
                text = extract_text(path)
                if len(text) < 100:
                    print(f"  ⚠️ Too short: {len(text)} chars.")
                    continue
                
                intel = self.parser.parse(text)
                raw_patterns = [pt.get("pattern", "") for pt in intel.get("patterns", [])]
                consolidated = self.merger.merge_list(raw_patterns, limit=7)
                notion_patterns = [{"name": pt[:100]} for pt in consolidated]
                score, _ = self.scorer.calculate_score(intel.get("patterns", []), set(), set(), intel)
                _, drive_link = self.gdrive.upload_file(path, GDRIVE_FOLDER_ID)
                
                profile = intel.get("candidate_profile", {})
                summary = profile.get("experience_summary", "") or intel.get("summary", "")
                sectors = profile.get("primary_sector", "Unclassified")
                if isinstance(sectors, str): sectors = [sectors]
                multi_sectors = [{"name": s} for s in sectors]

                props = {
                    "Primary Sector": {"multi_select": multi_sectors},
                    "Experience Patterns": {"multi_select": notion_patterns},
                    "Trajectory Grade": {"select": {"name": intel.get("career_path_quality", {}).get("trajectory_grade", "Neutral")}},
                    "v6.2 Score": {"number": score},
                    "구글드라이브 링크": {"url": drive_link}
                }
                if summary: props["경력 Summary"] = {"rich_text": [{"text": {"content": summary[:2000]}}]}
                
                requests.patch(f"https://api.notion.com/v1/pages/{nid}", headers=headers, json={"properties": props})
                print(f"  ✅ Recovered!")
                time.sleep(1)
            except Exception as e:
                print(f"  ❌ Error: {e}")

    def run(self):
        pages = self.get_live_pages()
        self.run_aggressive_dedupe(pages)
        self.run_fuzzy_recovery()
        print("\n✨ ALL PROCESSES COMPLETE.")

if __name__ == "__main__":
    UltimateRefiner().run()
