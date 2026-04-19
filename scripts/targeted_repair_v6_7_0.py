import os
import json
import time
import requests
import re
from docx import Document
import PyPDF2

# Project Imports
import sys
sys.path.append(os.getcwd())
from connectors.openai_api import OpenAIClient
from resume_parser import ResumeParser
from connectors.gdrive_api import GDriveConnector

# Configuration
DIR_RAW = r"C:\Users\cazam\Downloads\02_resume 전처리"
DIR_CONV = r"C:\Users\cazam\Downloads\02_resume_converted_docx"
GDRIVE_FOLDER_ID = "1VzJEeoXG239PVR3IoJM5jC28KccMdkth"

with open("secrets.json", "r") as f:
    secrets = json.load(f)

NOTION_TOKEN = secrets["NOTION_API_KEY"]
NOTION_DB_ID = secrets["NOTION_DATABASE_ID"]

headers = {
    "Authorization": f"Bearer {NOTION_TOKEN}",
    "Content-Type": "application/json",
    "Notion-Version": "2022-06-28"
}

TARGET_NAMES = ["이정한", "김정근", "심초아", "이종민", "장봉원", "전예찬", "전형준", "정예린", "정현구", "정현우", "정호진", "최재언", "황의영", "박주령", "강병수"]

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    text = []
    try:
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
        elif ext == '.docx':
            doc = Document(filepath)
            for p in doc.paragraphs:
                if p.text.strip(): text.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_content: text.append(" | ".join(dict.fromkeys(row_content)))
    except Exception:
        pass
    return "\n".join(text).strip()

def get_core_name(full_name):
    name = re.sub(r'\[.*?\]|\(.*?\)|\s|부문|_원본|_이력서|원본|_포트폴리오|지원서|이력서', '', full_name)
    name = re.sub(r'\d{2,4}', '', name)
    match = re.search(r'[가-힣]{2,4}', name)
    return match.group(0) if match else name.strip()

class TargetedRefiner:
    def __init__(self):
        self.openai_client = OpenAIClient(secrets["OPENAI_API_KEY"])
        self.parser = ResumeParser(self.openai_client)
        self.gdrive = GDriveConnector()
        self.file_cache = {}

    def preload_files(self):
        print("📂 Pre-loading all files...")
        for d in [DIR_RAW, DIR_CONV]:
            if not os.path.exists(d): continue
            for f in os.listdir(d):
                if f.lower().endswith(('.pdf', '.docx')) and not f.startswith("~$"):
                    path = os.path.join(d, f)
                    self.file_cache[path] = extract_text(path)[:2000] # Cache enough for name search
        print(f"  Indexed {len(self.file_cache)} files.")

    def find_file(self, name):
        core = get_core_name(name)
        # Priority 1: Filename match
        for path in self.file_cache:
            if core in os.path.basename(path):
                return path
        # Priority 2: Content match
        for path, preview in self.file_cache.items():
            if core in preview:
                return path
        return None

    def get_pages_to_fix(self):
        print("📋 Fetching pages to fix (Targeted + Blanks)...")
        url = f"https://api.notion.com/v1/databases/{NOTION_DB_ID}/query"
        pages = []
        has_more = True
        next_cursor = None
        while has_more:
            payload = {"page_size": 100}
            if next_cursor: payload["start_cursor"] = next_cursor
            resp = requests.post(url, headers=headers, json=payload).json()
            for p in resp.get('results', []):
                if p.get('archived'): continue
                props = p['properties']
                name_list = props.get('이름', {}).get('title', [])
                name = name_list[0]['plain_text'] if name_list else "Unknown"
                
                has_patterns = len(props.get('Experience Patterns', {}).get('multi_select', [])) > 0
                has_sector = len(props.get('Primary Sector', {}).get('multi_select', [])) > 0
                has_gdrive = props.get('구글드라이브 링크', {}).get('url') is not None
                
                is_target = any(tn in name for tn in TARGET_NAMES)
                is_blank = not has_patterns or not has_sector or not has_gdrive
                
                if is_target or is_blank:
                    pages.append({
                        "id": p['id'],
                        "name": name,
                        "need_patterns": not has_patterns,
                        "need_sector": not has_sector,
                        "need_gdrive": not has_gdrive
                    })
            has_more = resp.get('has_more', False)
            next_cursor = resp.get('next_cursor')
        return pages

    def fix(self, page):
        path = self.find_file(page['name'])
        if not path:
            print(f"  ❌ No file found for {page['name']}")
            return

        print(f"  🔄 Fixing {page['name']} using {os.path.basename(path)}")
        text = extract_text(path)
        
        try:
            intel = self.parser.parse(text)
            props = {}
            
            if page['need_patterns'] or any(tn in page['name'] for tn in TARGET_NAMES):
                raw_pts = [pt.get("pattern", "") for pt in intel.get("patterns", [])]
                props["Experience Patterns"] = {"multi_select": [{"name": pt[:100]} for pt in raw_pts[:7]]}
            
            if page['need_sector'] or any(tn in page['name'] for tn in TARGET_NAMES):
                profile = intel.get("candidate_profile", {})
                sectors = profile.get("primary_sector", "Unclassified")
                if isinstance(sectors, str): sectors = [sectors]
                props["Primary Sector"] = {"multi_select": [{"name": s} for s in sectors]}

            if page['need_gdrive'] or any(tn in page['name'] for tn in TARGET_NAMES):
                _, drive_link = self.gdrive.upload_file(path, GDRIVE_FOLDER_ID)
                props["구글드라이브 링크"] = {"url": drive_link}

            if props:
                requests.patch(f"https://api.notion.com/v1/pages/{page['id']}", headers=headers, json={"properties": props})
                print(f"  ✅ Fixed!")
            else:
                print(f"  ⚠️ No props to update.")
        except Exception as e:
            print(f"  ❌ Error: {e}")

    def run(self):
        pages = self.get_pages_to_fix()
        print(f"🚀 Found {len(pages)} pages to fix.")
        if not pages: return
        
        self.preload_files()
        
        for i, p in enumerate(pages):
            print(f"📝 [{i+1}/{len(pages)}] Processing: {p['name']}")
            self.fix(p)
            time.sleep(1)

if __name__ == "__main__":
    TargetedRefiner().run()
