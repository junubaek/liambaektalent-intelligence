import os
import sqlite3
import json
import time
import sys
import PyPDF2
import requests
from docx import Document
from concurrent.futures import ThreadPoolExecutor

sys.path.append(os.getcwd())
from resume_parser import ResumeParser
from headhunting_engine.normalization.pattern_merger import PatternMerger
from headhunting_engine.matching.scorer import Scorer
from connectors.openai_api import OpenAIClient
from connectors.gdrive_api import GDriveConnector

# Configuration
DIR_RAW = r"C:\Users\cazam\Downloads\02_resume 전처리"
DIR_CONV = r"C:\Users\cazam\Downloads\02_resume_converted_docx"
GDRIVE_FOLDER_ID = "1VzJEeoXG239PVR3IoJM5jC28KccMdkth"
NOTION_DB_ID = "31a22567-1b6f-8177-86da-ff626bb1e66c"
DB_PATH = "headhunting_engine/data/analytics.db"

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    text = []
    try:
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
        elif ext == '.docx':
            doc = Document(filepath)
            # 1. Extract Paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    text.append(p.text.strip())
            # 2. Extract Tables (Fix for NC 현종민 case)
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text.append(" | ".join(row_data))
    except Exception as e:
        print(f"  ⚠️ Error extracting {filepath}: {e}")
    return "\n".join(text).strip()

class UltimateRecovery:
    def __init__(self):
        with open("secrets.json", "r") as f:
            secrets = json.load(f)
        
        self.notion_token = secrets["NOTION_API_KEY"]
        self.headers = {
            "Authorization": f"Bearer {self.notion_token}",
            "Content-Type": "application/json",
            "Notion-Version": "2022-06-28"
        }
        self.openai_client = OpenAIClient(secrets["OPENAI_API_KEY"])
        self.parser = ResumeParser(self.openai_client)
        self.merger = PatternMerger()
        self.scorer = Scorer()
        self.gdrive = GDriveConnector()
        
        self.conn = sqlite3.connect(DB_PATH, check_same_thread=False)
        self.lock = False # Simple flag for DB safety in threads

    def get_notion_map(self):
        print("📋 Fetching existing Notion Hub map...")
        url = f"https://api.notion.com/v1/databases/{NOTION_DB_ID}/query"
        pages = []
        has_more = True
        start_cursor = None
        
        while has_more:
            payload = {"page_size": 100}
            if start_cursor: payload["start_cursor"] = start_cursor
            resp = requests.post(url, headers=self.headers, json=payload).json()
            pages.extend(resp.get('results', []))
            has_more = resp.get('has_more', False)
            start_cursor = resp.get('next_cursor')
            print(f"  -> Loaded {len(pages)} pages...")
            
        n_map = {}
        for p in pages:
            props = p.get('properties', {})
            title = props.get('이름', {}).get('title', [])
            if title:
                name = title[0].get('plain_text')
                n_map[name] = p['id']
        return n_map

    def process_candidate(self, name_base, filepath, notion_id):
        retries = 3
        for attempt in range(retries):
            try:
                print(f"🔄 Processing: {name_base} (Attempt {attempt+1})")
                
                # 1. Text Extraction
                text = extract_text(filepath)
                if not text or len(text) < 300:
                    print(f"  ⚠️ Skipping {name_base}: Text too short ({len(text) if text else 0} chars).")
                    # If it exists in Notion, update it with "Incomplete" tags
                    if notion_id:
                        props = {
                            "Primary Sector": {"select": {"name": "Incomplete_Profile"}},
                            "Experience Patterns": {"multi_select": []},
                            "v6.2 Score": {"number": 0}
                        }
                        requests.patch(f"https://api.notion.com/v1/pages/{notion_id}", headers=self.headers, json={"properties": props}, timeout=15)
                    return

                # 2. GDrive Link
                print(f"  📤 Ensuring GDrive link for {name_base}...")
                drive_id, drive_link = self.gdrive.upload_file(filepath, GDRIVE_FOLDER_ID)

                # 3. Deep Intelligence (v6.3.11 Fixed)
                print(f"  🧠 Deep Parsing (20k chars)...")
                intel = self.parser.parse(text)
                if not intel or intel.get("primary_sector") == "Insufficient_Data":
                    print(f"  ⚠️ Insufficient detail in {name_base}. Tagging as Incomplete.")
                    if notion_id:
                         props = {"Primary Sector": {"select": {"name": "Insufficient_Data"}}, "Experience Patterns": {"multi_select": []}}
                         requests.patch(f"https://api.notion.com/v1/pages/{notion_id}", headers=self.headers, json={"properties": props}, timeout=15)
                    return

                # 4. Pattern Merging
                raw_patterns = [p.get("pattern", "") for p in intel.get("patterns", [])]
                consolidated = self.merger.merge_list(raw_patterns, limit=7)
                notion_patterns = [{"name": p[:100]} for p in consolidated]

                # 5. Scoring
                score, _ = self.scorer.calculate_score(intel.get("patterns", []), set(), set(), intel)

                # 6. Notion Update
                profile = intel.get("candidate_profile", {})
                summary = profile.get("experience_summary", "")
                if not summary and "summary" in intel: summary = intel["summary"]
                
                props = {
                    "Primary Sector": {"select": {"name": profile.get("primary_sector", "Unclassified")}},
                    "Experience Patterns": {"multi_select": notion_patterns},
                    "Trajectory Grade": {"select": {"name": intel.get("career_path_quality", {}).get("trajectory_grade", "Neutral")}},
                    "v6.2 Score": {"number": score},
                    "구글드라이브 링크": {"url": drive_link}
                }
                if summary:
                    props["경력 Summary"] = {"rich_text": [{"text": {"content": summary[:2000]}}]}

                if notion_id:
                    requests.patch(f"https://api.notion.com/v1/pages/{notion_id}", headers=self.headers, json={"properties": props}, timeout=15)
                    print(f"  ✅ Updated Notion Page: {name_base}")
                else:
                    # Final safety check: Does this page exist by title?
                    props["이름"] = {"title": [{"text": {"content": name_base}}]}
                    requests.post("https://api.notion.com/v1/pages", headers=self.headers, json={"parent": {"database_id": NOTION_DB_ID}, "properties": props}, timeout=15)
                    print(f"  ✅ Created New Notion Page: {name_base}")
                
                return # Success!

            except Exception as e:
                print(f"  ⚠️ Error in {name_base} (Attempt {attempt+1}): {e}")
                time.sleep(2 ** attempt) # Exponential backoff
        
        print(f"  ❌ Failed {name_base} after {retries} retries.")

    def run(self):
        notion_map = self.get_notion_map()
        
        # Collect Files & Deduplicate
        all_files = {}
        for d in [DIR_RAW, DIR_CONV]:
            if not os.path.exists(d): continue
            for f in os.listdir(d):
                if f.lower().endswith(('.pdf', '.docx')) and not f.startswith("~$"):
                    name = os.path.splitext(f)[0]
                    # Deduplicate: if multiple files exist, take the largest or newest?
                    # Here we just take the first we see.
                    if name not in all_files:
                        all_files[name] = os.path.join(d, f)
        
        print(f"📂 Found {len(all_files)} candidates to process.")
        
        # Sequential processing for 100% stability in recovery
        for i, (name, path) in enumerate(all_files.items()):
            nid = notion_map.get(name)
            self.process_candidate(name, path, nid)
            
            if (i+1) % 10 == 0:
                print(f"🏁 Checkpoint: {i+1}/{len(all_files)} processed.")
            
            time.sleep(0.5)

if __name__ == "__main__":
    recovery = UltimateRecovery()
    recovery.run()
