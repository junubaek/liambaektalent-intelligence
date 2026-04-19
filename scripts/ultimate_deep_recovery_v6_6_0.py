import os
import json
import time
import requests
import re
from docx import Document
import PyPDF2
from concurrent.futures import ThreadPoolExecutor

# Project Imports
import sys
sys.path.append(os.getcwd())
from connectors.openai_api import OpenAIClient
from resume_parser import ResumeParser

# Configuration
DIR_RAW = r"C:\Users\cazam\Downloads\02_resume 전처리"
DIR_CONV = r"C:\Users\cazam\Downloads\02_resume_converted_docx"

with open("secrets.json", "r") as f:
    secrets = json.load(f)

NOTION_TOKEN = secrets["NOTION_API_KEY"]
NOTION_DB_ID = secrets["NOTION_DATABASE_ID"]

headers = {
    "Authorization": f"Bearer {NOTION_TOKEN}",
    "Content-Type": "application/json",
    "Notion-Version": "2022-06-28"
}

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    text = []
    try:
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text.append(page.extract_text() or "")
        elif ext == '.docx':
            doc = Document(filepath)
            for p in doc.paragraphs:
                if p.text.strip(): text.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_content = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_content: text.append(" | ".join(dict.fromkeys(row_content)))
    except Exception as e:
        pass
    return "\n".join(text).strip()

def get_core_name(full_name):
    name = re.sub(r'\[.*?\]|\(.*?\)|\s|부문|_원본|_이력서|원본|_포트폴리오|지원서|이력서', '', full_name)
    name = re.sub(r'\d{2,4}', '', name)
    match = re.search(r'[가-힣]{2,4}', name)
    return match.group(0) if match else name.strip()

class DeepRefiner:
    def __init__(self):
        self.openai_client = OpenAIClient(secrets["OPENAI_API_KEY"])
        self.parser = ResumeParser(self.openai_client)
        self.file_cache = {}

    def preload_files(self):
        print("📂 Pre-loading and indexing all files (Deep Search Mode)...")
        for d in [DIR_RAW, DIR_CONV]:
            if not os.path.exists(d): continue
            for f in os.listdir(d):
                if f.lower().endswith(('.pdf', '.docx')) and not f.startswith("~$"):
                    path = os.path.join(d, f)
                    # We store path and first 500 chars for fast name searching
                    text_preview = extract_text(path)[:1000]
                    self.file_cache[path] = text_preview
        print(f"  Indexed {len(self.file_cache)} files.")

    def find_file_by_content(self, name):
        # Try exact filename match first
        core = get_core_name(name)
        for path in self.file_cache:
            if core in os.path.basename(path):
                return path
        # Deep Search: Look inside content
        for path, preview in self.file_cache.items():
            if core in preview:
                return path
        return None

    def get_blank_pages(self):
        print("📋 Identifying pages with missing Summary or Patterns...")
        url = f"https://api.notion.com/v1/databases/{NOTION_DB_ID}/query"
        blanks = []
        has_more = True
        next_cursor = None
        while has_more:
            payload = {"page_size": 100}
            if next_cursor: payload["start_cursor"] = next_cursor
            resp = requests.post(url, headers=headers, json=payload).json()
            for p in resp.get('results', []):
                if p.get('archived'): continue
                props = p['properties']
                has_patterns = len(props.get('Experience Patterns', {}).get('multi_select', [])) > 0
                has_summary = len(props.get('경력 Summary', {}).get('rich_text', [])) > 0
                
                if not has_patterns or not has_summary:
                    name_list = props.get('이름', {}).get('title', [])
                    name = name_list[0]['plain_text'] if name_list else "Unknown"
                    blanks.append({"id": p['id'], "name": name, "need_patterns": not has_patterns, "need_summary": not has_summary})
            has_more = resp.get('has_more', False)
            next_cursor = resp.get('next_cursor')
        return blanks

    def recover(self, target):
        path = self.find_file_by_content(target['name'])
        if not path:
            print(f"  ❌ No file found for {target['name']}")
            return

        print(f"  🔄 Recovering {target['name']} using {os.path.basename(path)}")
        text = extract_text(path)
        if len(text) < 50: # Relaxed for summary
            print(f"  ⚠️ Text too short ({len(text)} chars)")
            return

        try:
            # Re-parse or just get summary if patterns exist
            intel = self.parser.parse(text)
            props = {}
            
            if target['need_patterns']:
                raw_pts = [pt.get("pattern", "") for pt in intel.get("patterns", [])]
                props["Experience Patterns"] = {"multi_select": [{"name": pt[:100]} for pt in raw_pts[:7]]}
            
            if target['need_summary']:
                summary = intel.get("candidate_profile", {}).get("experience_summary", "") or intel.get("summary", "")
                if not summary and len(text) > 100:
                    # Forced summarization if AI missed it
                    summary = text[:500].replace("\n", " ") + "..."
                if summary:
                    props["경력 Summary"] = {"rich_text": [{"text": {"content": summary[:2000]}}]}

            if props:
                requests.patch(f"https://api.notion.com/v1/pages/{target['id']}", headers=headers, json={"properties": props})
                print(f"  ✅ Success!")
            else:
                print(f"  ⚠️ No new data extracted.")
        except Exception as e:
            print(f"  ❌ AI Error: {e}")

    def run(self):
        blanks = self.get_blank_pages()
        print(f"🚀 Found {len(blanks)} potential blanks to fill.")
        if not blanks: return
        
        self.preload_files()
        
        for i, b in enumerate(blanks):
            print(f"📝 [{i+1}/{len(blanks)}] Processing: {b['name']}")
            self.recover(b)
            time.sleep(1)

if __name__ == "__main__":
    DeepRefiner().run()
