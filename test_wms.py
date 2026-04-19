
import os
import docx

dirs = [r'C:\Users\cazam\Downloads\02_resume_converted_v8']
keywords = ['WMS', '이사만루', '채용팀장']
docs = []

for d in dirs:
    if os.path.exists(d):
        for f in os.listdir(d):
            if any(k in f for k in keywords) and f.endswith('.docx'):
                docs.append(os.path.join(d, f))

for path in docs:
    try:
        doc = docx.Document(path)
        print(f'=== {os.path.basename(path)} ===')
        text = ''
        for para in doc.paragraphs[:10]:
            if para.text.strip(): text += para.text + '\n'
        for table in doc.tables[:2]:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip(): text += cell.text + ' '
        print(text[:200].replace('\n', ' '))
    except Exception as e:
        print('Error:', e)

