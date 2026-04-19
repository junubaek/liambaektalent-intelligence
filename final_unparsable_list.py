import os
import pdfplumber
from docx import Document
import fitz
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed

TARGET_DIRS = [
    r"C:\Users\cazam\Downloads\02_resume 전처리",
    r"C:\Users\cazam\Downloads\02_resume_converted_v8"
]

def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    text = ""
    try:
        if ext == "pdf":
            with fitz.open(filepath) as doc:
                for page in doc: text += page.get_text()
        elif ext == "docx":
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            for t in doc.tables:
                for r in t.rows:
                    for ce in r.cells:
                        if ce.text: text += ce.text + " "
    except Exception:
        pass
    return text.strip()

def check_file(filepath):
    filename = os.path.basename(filepath)
    ext = filename.rsplit(".", 1)[-1].lower()
    
    # We ignore raw .doc since we assume it's converted
    if ext == "doc":
        return None
        
    text = extract_text(filepath)
    # Check if text extraction yielded less than 50 characters
    if len(text) < 50:
        return filename
    return None

def main():
    files = []
    for d in TARGET_DIRS:
        if os.path.exists(d):
            files.extend([os.path.join(d, f) for f in os.listdir(d) 
                         if f.endswith(('.pdf', '.docx')) and not f.startswith('~$')])
            
    print(f"Total valid PDF/DOCX files found: {len(files)}")
    
    failed_files = []
    
    with ThreadPoolExecutor(max_workers=8) as executor:
        futures = [executor.submit(check_file, fp) for fp in files]
        for future in tqdm(as_completed(futures), total=len(files)):
            res = future.result()
            if res:
                failed_files.append(res)
                
    failed_files = sorted(list(set(failed_files)))
    
    arti_path = r"C:\Users\cazam\.gemini\antigravity\brain\40516708-d2c8-4e50-a5f8-ff10183de737\artifacts\unparsable_resumes.md"
    doc = f"# 📄 완전히 파싱 불가능한 스캔/이미지 이력서 리스트\n\n"
    doc += f"- 순수 PDF 및 DOCX 파일만 대상 (구형 .doc는 목록에서 완전히 제외)\n"
    doc += f"- 총 점검 대상 수량: {len(files)}개\n"
    doc += f"- **최종 파싱 불가능 수량:** **{len(failed_files)}개**\n"
    doc += "*(대부분 글자가 아닌 이미지 기반의 스캔 PDF 파일이거나 암호화된 문서입니다.)*\n\n"
    doc += "| 연번 | 파일명 |\n|---|---|\n"
    
    for i, f in enumerate(failed_files):
        doc += f"| {i+1} | {f} |\n"
        
    with open(arti_path, 'w', encoding='utf-8') as f:
        f.write(doc)
    print(f"\nFound {len(failed_files)} totally unparsable files.")

if __name__ == "__main__":
    main()
