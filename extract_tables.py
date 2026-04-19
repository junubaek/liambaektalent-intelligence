import sqlite3
import json
import os
import sys
import docx
import time

sys.stdout.reconfigure(encoding='utf-8')

def extract_docx_full(filepath):
    try:
        doc = docx.Document(filepath)
        text_blocks = []
        
        # 1. 문단 텍스트 추출
        for p in doc.paragraphs:
            if p.text.strip():
                text_blocks.append(p.text.strip())
                
        # 2. 표 내부 텍스트 추출 (핵심)
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        # 여러 줄이 있을 수 있으므로 공백으로 치환
                        row_text.append(cell.text.strip().replace('\n', ' '))
                if row_text:
                    text_blocks.append(" | ".join(row_text))
                    
        return "\n".join(text_blocks)
    except Exception as e:
        print(f"Failed to read docx {filepath}: {e}")
        return ""

def run_table_extraction():
    print("=== Table Data Extraction & DB Update ===")
    
    with open('processed_step6.json', 'r', encoding='utf-8') as f:
        d = json.load(f)
        
    zero_uids = [uid for uid, v in d.items() if v.get('edge_count', 0) == 0]
    
    db = sqlite3.connect('candidates.db')
    c = db.cursor()
    
    names_map = {}
    for uid in zero_uids:
        c.execute('SELECT name_kr FROM candidates WHERE id=?', (uid,))
        row = c.fetchone()
        if row:
            names_map[row[0]] = uid
            
    folder = r'C:\Users\cazam\Downloads\02_resume_converted_v8'
    files = os.listdir(folder) if os.path.exists(folder) else []
    
    updated_count = 0
    uids_to_reprocess = []
    
    for name, uid in names_map.items():
        for f in files:
            if name in f and f.endswith('.docx'):
                filepath = os.path.join(folder, f)
                text = extract_docx_full(filepath)
                if len(text) > 300: 
                    # Update SQLite
                    c.execute("UPDATE candidates SET raw_text=? WHERE id=?", (text, uid))
                    db.commit()
                    updated_count += 1
                    uids_to_reprocess.append(uid)
                break
                
    print(f"-> 표 추출 포함 DB 업데이트 완료: {updated_count}명")
    
    # Remove from processed_step6.json so parser will run them again
    for uid in uids_to_reprocess:
        if uid in d:
            del d[uid]
            
    with open('processed_step6.json', 'w', encoding='utf-8') as f:
        json.dump(d, f, ensure_ascii=False, indent=2)
        
    print("-> processed_step6.json 초기화 완료. 이제 dynamic_parser_step6.py를 다시 구동합니다.")
    
    # 3명만 샘플로 추출된 텍스트 보여주기
    print("\n--- [샘플 추출 텍스트 (앞 300자)] ---")
    c.execute("SELECT name_kr, raw_text FROM candidates WHERE id IN ({seq}) LIMIT 3".format(seq=','.join(['?']*len(uids_to_reprocess))), uids_to_reprocess)
    for row in c.fetchall():
        print(f"\n[{row[0]}]\n{row[1][:300]}...")

if __name__ == "__main__":
    run_table_extraction()
