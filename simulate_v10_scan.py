import os
import sqlite3
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm
import pdfplumber
from docx import Document
import fitz

TARGET_DIRS = [
    r"C:\Users\cazam\Downloads\02_resume 전처리",
    r"C:\Users\cazam\Downloads\02_resume_converted_v8"
]
DB_PATH = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\candidates.db"

def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    text = ""
    try:
        if ext == "pdf":
            with fitz.open(filepath) as doc:
                for page in doc: text += page.get_text()
        elif ext in ("docx", "doc"):
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            for t in doc.tables:
                for r in t.rows:
                    for ce in r.cells:
                        if ce.text: text += ce.text + " "
    except Exception:
        pass
    return text.strip()

def check_file(filepath):
    text = extract_text(filepath)
    if len(text) < 50:
        return (False, filepath, "too_short")
    
    doc_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    return (True, filepath, doc_hash)

def main():
    print("Collecting files...")
    files = []
    for d in TARGET_DIRS:
        if os.path.exists(d):
            files.extend([os.path.join(d, f) for f in os.listdir(d) 
                         if f.endswith(('.pdf', '.docx', '.doc')) and not f.startswith('~$')])
            
    print(f"Total files found: {len(files)}")
    
    # Load all hashes from DB into memory for O(1) checking
    print("Loading SQLite DB hashes...")
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    db_hashes = set(row[0] for row in c.execute('SELECT document_hash FROM candidates WHERE document_hash IS NOT NULL').fetchall())
    conn.close()
    print(f"Loaded {len(db_hashes)} existing hashes from DB.")
    
    existing_count = 0
    new_files = []
    error_count = 0

    print("Hashing and analyzing files...")
    with ThreadPoolExecutor(max_workers=8) as executor:
        futures = [executor.submit(check_file, fp) for fp in files]
        for future in tqdm(as_completed(futures), total=len(files)):
            try:
                ok, filepath, res = future.result()
                if not ok:
                    error_count += 1
                    continue
                
                doc_hash = res
                if doc_hash in db_hashes:
                    existing_count += 1
                else:
                    new_files.append(os.path.basename(filepath))
            except Exception as e:
                error_count += 1

    print("\n" + "="*40)
    print("📋 [사전 스캔 분석 결과]")
    print(f"- 기존 파일 (스킵 예정): {existing_count}개")
    print(f"- 신규 파일 (처리 예정): {len(new_files)}개")
    print(f"- 추출 실패/오류 (Too short 등): {error_count}개")
    print(f"- 총계: {len(files)}개")
    print("="*40)
    
    print("\n🚀 [신규 파일 목록 상위 20개]")
    for i, f in enumerate(new_files[:20]):
        print(f"{i+1:02d}. {f}")

if __name__ == "__main__":
    main()
