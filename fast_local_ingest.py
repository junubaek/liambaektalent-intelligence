import os
import json
import time
import uuid
import sqlite3
import hashlib
import re
from datetime import datetime
import fitz  # PyMuPDF
from docx import Document
import google.generativeai as genai
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed

# Target Folder
TARGET_DIR = r"C:\Users\cazam\Downloads\02_resume 전처리"
DB_PATH = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\candidates.db"
SECRETS_PATH = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템\secrets.json"

with open(SECRETS_PATH, "r", encoding="utf-8") as f:
    secrets = json.load(f)

# Config Gemini
genai.configure(api_key=secrets["GEMINI_API_KEY"])
model = genai.GenerativeModel("gemini-2.0-flash-exp") # Using stable/fast model

MEGA_PROMPT = """이력서 텍스트를 분석해서 아래 구조를 JSON Object 형식으로 출력해. 반드시 코드블럭 없이 JSON만 반환할 것.

{{
  "name_kr": "본문에서 순수 본명만 추출 (직무 꼬리표 붙이지 말 것). 못 찾으면 null",
  "phone": "010-XXXX-XXXX 형태 숫자만. 못 찾으면 null",
  "email": "이메일주소. 못 찾으면 null",
  "birth_year": 출생연도 숫자 4자리 (예: 1990). 못 찾으면 null,
  "summary": "주요 경력 2줄 요약 (개인정보 절대 포함 금지!!!)",
  "sector": "가장 관련 깊은 도메인/산업분야 분류 (예: SW, Finance, Marketing 등)",
  "education_json": [
    {{"school": "대학교 이상 학교명", "major": "전공", "degree": "학사/석사 등", "year": "졸업년도"}}
  ],
  "careers_json": [
    {{
      "company": "회사명",
      "title": "직책/부서",
      "start_date": "YYYY.MM",
      "end_date": "YYYY.MM (또는 현재)"
    }}
  ]
}}

이력서 본문:
{text}
"""

def extract_text(filepath):
    ext = filepath.rsplit(".", 1)[-1].lower()
    text = ""
    try:
        if ext == "pdf":
            with fitz.open(filepath) as doc:
                for page in doc: text += page.get_text()
        elif ext in ("docx", "doc"):
            doc = Document(filepath)
            text = "\n".join(p.text for p in doc.paragraphs)
            for t in doc.tables:
                for r in t.rows:
                    for ce in r.cells:
                        if ce.text: text += ce.text + " "
    except Exception as e:
        print(f"Extraction error for {filepath}: {e}")
    return text.strip()

def calculate_career_stats(careers):
    if not careers: return "미상", 0.0
    total_months = 0
    latest_end_dt = datetime(1900, 1, 1)
    curr_comp = "미상"
    
    def parse_dt(d_str):
        if not d_str or '현재' in d_str or '재직' in d_str or 'ing' in d_str.lower(): return datetime.now()
        m = re.findall(r'(\d{4})[^\d]*(\d{1,2})', d_str)
        if m: return datetime(int(m[0][0]), int(m[0][1]), 1)
        m2 = re.findall(r'(\d{4})', d_str)
        if m2: return datetime(int(m2[0]), 1, 1)
        return datetime.now()

    for c in careers:
        st = parse_dt(c.get('start_date', ''))
        ed = parse_dt(c.get('end_date', ''))
        if ed < st: ed = st
        months = (ed.year - st.year) * 12 + ed.month - st.month
        total_months += max(months, 0)
        
        if ed >= latest_end_dt:
            latest_end_dt = ed
            if c.get('company'): curr_comp = c.get('company')
            
    total_y = round(total_months / 12.0, 1)
    return curr_comp, total_y

def process_file(filepath):
    filename = os.path.basename(filepath)
    text = extract_text(filepath)
    if len(text) < 50:
        return False, "Not enough text"
        
    doc_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    
    # MD5 Hash Duplicate Check
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('SELECT COUNT(*) FROM candidates WHERE document_hash=?', (doc_hash,))
    if c.fetchone()[0] > 0:
        conn.close()
        return False, "Already processed (Hash Match)"
    conn.close()

    # Gemini Parsing
    parsed = None
    for _ in range(2):
        try:
            prompt = MEGA_PROMPT.format(text=f"[파일명: {filename}]\n\n" + text[:8000])
            res = model.generate_content(prompt)
            raw = res.text.replace("```json", "").replace("```", "").strip()
            parsed = json.loads(raw)
            break
        except Exception:
            time.sleep(1)
            
    if not parsed:
        return False, "Gemini Parsing Failed"

    name_kr = parsed.get("name_kr") or filename[:10]
    email = parsed.get("email", "")
    phone = parsed.get("phone", "")
    birth_year = parsed.get("birth_year", 0)
    sector = parsed.get("sector", "미분류")
    summary = parsed.get("summary", "")
    careers = parsed.get("careers_json", [])
    edu = parsed.get("education_json", [])
    
    current_company, total_years = calculate_career_stats(careers)

    # Insert into SQLite
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c_id = str(uuid.uuid4())
    try:
        c.execute('''
            INSERT INTO candidates (
                id, name_kr, email, phone, birth_year, sector, profile_summary, total_years, current_company,
                raw_text, document_hash, is_duplicate, is_parsed, created_at, updated_at, 
                careers_json, education_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, 1, datetime('now'), datetime('now'), ?, ?)
        ''', (c_id, name_kr, email, phone, birth_year, sector, summary, total_years, current_company,
              text, doc_hash, json.dumps(careers, ensure_ascii=False), json.dumps(edu, ensure_ascii=False)))
        conn.commit()
    except Exception as e:
        conn.close()
        return False, f"DB Error: {e}"
    
    conn.close()
    return True, f"{name_kr} ({current_company}, {total_years} yrs)"

def main():
    if not os.path.exists(TARGET_DIR):
        print(f"Directory not found: {TARGET_DIR}")
        return

    files = [os.path.join(TARGET_DIR, f) for f in os.listdir(TARGET_DIR) 
             if f.lower().endswith(('.pdf', '.docx', '.doc')) and not f.startswith('~$')]
    
    print(f"Total target files: {len(files)}")
    
    success = 0
    skipped = 0
    
    # Use 5 workers for faster processing
    with ThreadPoolExecutor(max_workers=5) as executor:
        futures = {executor.submit(process_file, fp): os.path.basename(fp) for fp in files}
        for future in tqdm(as_completed(futures), total=len(files)):
            filename = futures[future]
            try:
                ok, msg = future.result()
                if ok:
                    success += 1
                else:
                    if "Hash Match" not in msg:
                        print(f"\n[ERR] {filename}: {msg}")
                    skipped += 1
            except Exception as e:
                skipped += 1
                print(f"\n[EXC] {filename}: {e}")
                
    print(f"\nCompleted. Success: {success}, Skipped (Already in DB/Error): {skipped}")

if __name__ == "__main__":
    main()
