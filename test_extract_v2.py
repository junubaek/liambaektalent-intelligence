
from docx import Document
import os

def extract_all_text(path):
    doc = Document(path)
    text_parts = []
    
    # Extract from paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text_parts.append(p.text.strip())
            
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                # Cells can have multiple paragraphs
                cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))
                
    return "\n".join(text_parts)

path = r'C:\Users\cazam\Downloads\02_resume_converted_v8\김시형(총무팀 자산관리)부문_원본.docx'
if os.path.exists(path):
    text = extract_all_text(path)
    with open('test_extract_v2.txt', 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"Extraction V2 successful. Length: {len(text)}")
else:
    print("File not found")
