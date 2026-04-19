import sqlite3
import json
import docx
import google.generativeai as genai

with open('secrets.json', 'r') as f: secrets = json.load(f)
genai.configure(api_key=secrets.get('GEMINI_API_KEY', ''))
model = genai.GenerativeModel('gemini-2.5-flash-lite', generation_config={'response_mime_type': 'application/json'})

path = r'C:\Users\cazam\Downloads\02_resume_converted_v8\강민아(아시아사업팀)부문_원본.docx'
doc = docx.Document(path)

full_text = []
for para in doc.paragraphs: 
    full_text.append(para.text)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells: 
            full_text.append(cell.text)
raw_text = '\n'.join(full_text)

prompt = f"""아래 이력서 전문에서 직무(careers) 배열을 추출하세요.

[JSON 출력 포맷]
{{
  "careers": [
    {{ "company": "회사명(없으면 미상)", "title": "직급/직무", "start_date": "YYYY-MM 또는 미상", "end_date": "미상/현재" }}
  ]
}}

[이력서 전문]
{raw_text[:13000]}"""

res = model.generate_content(prompt).text.strip()
print('--- GEMINI EXTRACTION RESULT ---')
print(res)

conn = sqlite3.connect(r'c:\Users\cazam\Downloads\이력서자동분석검색시스템\candidates.db')
c = conn.cursor()
c.execute("UPDATE candidates SET raw_text=?, careers_json=? WHERE name_kr='강민아'", (raw_text, res))
conn.commit()
conn.close()
