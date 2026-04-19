import os
import sqlite3
import json
import time
import sys
import PyPDF2
import hashlib
import threading
from docx import Document
from concurrent.futures import ThreadPoolExecutor

# Define base project path for absolute imports
PROJECT_ROOT = r"C:\Users\cazam\Downloads\이력서자동분석검색시스템"
sys.path.append(PROJECT_ROOT)

from resume_parser import ResumeParser
from connectors.gemini_api import GeminiClient
from connectors.notion_api import NotionClient
from connectors.gdrive_api import GDriveConnector
from connectors.openai_api import OpenAIClient
from connectors.pinecone_api import PineconeClient

# Configuration v8.0
DIR_RAW = r"C:\Users\cazam\Downloads\02_resume 전처리"
DIR_CONV = r"C:\Users\cazam\Downloads\02_resume_converted_docx"
DIR_CONV_V8 = r"C:\Users\cazam\Downloads\02_resume_converted_v8"
GDRIVE_FOLDER_ID = "1jAf8VrxEaVBScwscLANWAuSJG1YvfFqz" 

# Canonical Main Sectors for Normalization (v2 - Deduplicated)
CANONICAL_MAIN_SECTORS = sorted(list(set([
    "영업 (Sales)", "마케팅 (Marketing)", "HR (Human Resources)", "총무 (General Affairs)",
    "Finance (재무/회계)", "STRATEGY (전략)", "디자인 (Design)", "법무 (Legal)",
    "물류/유통 (Logistics & SCM)", "MD (Commerce)", "PRODUCT (제품 기획)", "DATA (데이터)",
    "SW (Software)", "HW (Hardware)", "반도체 (Semiconductor)", "AI (Artificial Intelligence)",
    "보안 (Security)", "기타"
])))

# Canonical Sub-Sectors for Normalization (v3 - Strict Unique)
_SUB_SET = [
    "해외영업", "B2B영업", "기술영업(Pre-sales)", "영업지원", "영업기획",
    "퍼포먼스", "그로스", "브랜드", "콘텐츠(인플루언서 협업/제휴 포함)", "언론홍보(PR)", "마케팅기획",
    "채용(TA)", "평가보상(C&B)", "급여(Payroll)", "노무(ER)", "인사기획", "교육(L&D)",
    "복리후생 운영", "자산관리(부동산; 법인자산; IT 비품 구매 및 관리)",
    "재무회계", "자금", "세무", "IR", "M&A", "내부통제_감사", "FP&A(경영분석)", "회계사(CPA)", "투자담당자(Investment/VC/PE)",
    "전략_경영기획", "Business Operation(프로세스 효율화)", "신사업 발굴 및 런칭",
    "UIUX", "제품", "웹", "디자인 기획 및 시스템 구축",
    "일반법무", "컴플라이언스", "지적재산권(IP)", "변호사",
    "구매", "SCM(수요예측/공급망)", "유통망 관리", "물류기획 및 프로세스 최적화",
    "상품기획(Selection)", "소싱MD(해외 직소싱/원가)", "영업MD(채널 매출 관리)",
    "Product Owner(PO)", "Project Manager(PM)", "서비스기획(화면설계)", "TPM",
    "데이터분석가", "데이터엔지니어", "데이터사이언티스트", "DBA(DBMS 운영; 성능 최적화; 보안)",
    "BE(서버)", "FE(웹)", "DevOps_SRE", "인프라_Cloud", "Mobile(iOS; Android)",
    "회로설계(PCB)", "기구설계", "로보틱스", "자동화(PLC)", "임베디드", "FAE_CE",
    "HW(SoC; RTL)", "SW(하단 드라이버)", "FW(컨트롤러)", "공정(Yield/FAB)",
    "엔지니어(Serving/MLOps)", "리서쳐(모델링)", "기획(AI Governance; DT; AT; AX 전략 설계)",
    "정보보안(인프라/인증)", "개인정보보호(CPO; 규제 대응)",
    "기타"
]
CANONICAL_SUB_SECTORS = sorted(list(set(_SUB_SET)))

# Thread-local storage for clients
thread_local = threading.local()

def get_clients():
    if not hasattr(thread_local, "clients"):
        secrets_path = os.path.join(PROJECT_ROOT, "secrets.json")
        with open(secrets_path, "r") as f:
            secrets = json.load(f)
        
        gemini_client = GeminiClient(secrets["GEMINI_API_KEY"])
        thread_local.clients = {
            "parser": ResumeParser(gemini_client),
            "notion": NotionClient(secrets["NOTION_API_KEY"]),
            "gdrive": GDriveConnector(secrets_path=secrets_path),
            "openai": OpenAIClient(secrets["OPENAI_API_KEY"]),
            "pinecone": None # Initialized below
        }
        
        pc_host = secrets.get("PINECONE_HOST", "").rstrip("/")
        if pc_host and not pc_host.startswith("https://"): pc_host = f"https://{pc_host}"
        thread_local.clients["pinecone"] = PineconeClient(secrets["PINECONE_API_KEY"], pc_host)
        
        thread_local.notion_db_id = secrets.get("NOTION_DATABASE_ID", "31f22567-1b6f-8190-b3a8-dc4f8422f01b")
        
    return thread_local.clients, thread_local.notion_db_id

def extract_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif ext == '.docx':
            doc = Document(filepath)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            text = "\n".join(text_parts)
    except Exception as e:
        print(f"Error extracting {filepath}: {e}")
    return text.strip()

def is_meaningless(text, filepath):
    if not text or len(text) < 150: return True
    if os.path.basename(filepath).startswith("~$"): return True
    return False

def process_candidate(name_base, file_map):
    try:
        clients, notion_db_id = get_clients()
        filepath = file_map[name_base]
        text = extract_text(filepath)
        if is_meaningless(text, filepath): return False

        print(f"[*] Processing: {name_base}")
        
        # A. GDrive
        drive_id, drive_link = clients["gdrive"].upload_file(filepath, GDRIVE_FOLDER_ID)

        # B. Parse
        parsed_data = clients["parser"].parse(text)
        if not parsed_data: return False

        # C. Notion Sync
        profile = parsed_data.get("candidate_profile", {})
        patterns = parsed_data.get("patterns", [])
        
        main_sectors = [{"name": s.strip()[:100]} for s in profile.get("main_sectors", []) if s.strip() in CANONICAL_MAIN_SECTORS]
        sub_sectors = [{"name": s.strip()[:100]} for s in profile.get("sub_sectors", []) if s.strip() in CANONICAL_SUB_SECTORS]
        
        exp_summary = profile.get("experience_summary", "")
        if isinstance(exp_summary, list): exp_summary = "\n".join(exp_summary)
        
        props = {
            "이름": {"title": [{"text": {"content": name_base}}]},
            "Main Sectors": {"multi_select": main_sectors},
            "Sub Sectors": {"multi_select": sub_sectors},
            "Seniority Bucket": {"select": {"name": (profile.get("seniority_bucket") or "Unknown")}},
            "Functional Patterns": {"multi_select": [{"name": p.get("verb_object", "").strip()[:100]} for p in patterns if p.get("verb_object", "").strip()]},
            "Context Tags": {"rich_text": [{"text": {"content": ", ".join(profile.get("context_tags", []))[:2000]}}]},
            "Experience Summary": {"rich_text": [{"text": {"content": exp_summary[:2000]}}]},
            "AI Last Optimized": {"date": {"start": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())}},
            "구글드라이브 링크": {"url": drive_link}
        }

        filter_name = {"property": "이름", "title": {"equals": name_base}}
        existing = clients["notion"].query_database(notion_db_id, limit=1, filter_criteria=filter_name)
        
        if existing.get('results'):
            clients["notion"].update_page_properties(existing['results'][0]['id'], props)
        else:
            clients["notion"].create_page(notion_db_id, props)
        
        # D. Vector Upsert
        vec_text = f"Candidate: {name_base} | Sectors: {', '.join(profile.get('main_sectors', []))} | Summary: {exp_summary}"
        embedding = clients["openai"].embed_content(vec_text)
        if embedding:
            vector_id = hashlib.md5(name_base.encode('utf-8')).hexdigest()
            clients["pinecone"].upsert([{"id": vector_id, "values": embedding, "metadata": {"name": name_base}}])
        
        return True
    except Exception as e:
        print(f"  ❌ Error for {name_base}: {e}")
        return False

def unified_process_v8():
    print("🚀 Starting AI Talent Intelligence OS v8.0 Unified Sync...")
    
    # 1. File Mapping
    file_map = {}
    if os.path.exists(DIR_RAW):
        for f in os.listdir(DIR_RAW):
            if f.lower().endswith(('.pdf', '.docx')) and not f.startswith("~$"):
                name_base = os.path.splitext(f)[0]
                file_map[name_base] = os.path.join(DIR_RAW, f)
            
    if os.path.exists(DIR_CONV):
        for f in os.listdir(DIR_CONV):
            if f.lower().endswith(('.docx')) and not f.startswith("~$"):
                name_base = os.path.splitext(f)[0]
                file_map[name_base] = os.path.join(DIR_CONV, f)

    if os.path.exists(DIR_CONV_V8):
        for f in os.listdir(DIR_CONV_V8):
            if f.lower().endswith(('.docx')) and not f.startswith("~$"):
                name_base = os.path.splitext(f)[0]
                file_map[name_base] = os.path.join(DIR_CONV_V8, f)
            
    all_names = list(file_map.keys())
    print(f"💡 Total candidates found: {len(all_names)}")

    # PILOT MODE (Comment out for full run)
    # all_names = all_names[:3]

    success_count = 0
    max_workers = 5
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        from functools import partial
        results = list(executor.map(partial(process_candidate, file_map=file_map), all_names))
        success_count = sum(1 for r in results if r)

    print(f"\n✨ v8.0 Sync Complete. Success: {success_count}/{len(all_names)}.")

if __name__ == "__main__":
    unified_process_v8()
