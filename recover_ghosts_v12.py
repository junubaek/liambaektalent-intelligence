import os
import sqlite3
import json
import docx
import google.generativeai as genai
from concurrent.futures import ThreadPoolExecutor, as_completed
from tqdm import tqdm

with open('secrets.json', 'r', encoding='utf-8') as f: secrets = json.load(f)
genai.configure(api_key=secrets.get('GEMINI_API_KEY', ''))
model = genai.GenerativeModel('gemini-2.5-flash-lite', generation_config={'response_mime_type': 'application/json'})

DIR_V8 = r'C:\Users\cazam\Downloads\02_resume_converted_v8'
DIR_RAW = r'C:\Users\cazam\Downloads\02_resume 전처리'

def parse_docx(path):
    try:
        doc = docx.Document(path)
        full_text = []
        for para in doc.paragraphs: 
            if para.text.strip(): full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells: 
                    if cell.text.strip(): full_text.append(cell.text)
        return '\n'.join(full_text)
    except:
        return ""

def process_file(task_info):
    db_id, name_kr, file_path = task_info
    raw_text = parse_docx(file_path)
    
    if len(raw_text) < 100:
        return db_id, name_kr, raw_text, None, None

    prompt = f"""아래 이력서 전문에서 직무(careers)를 추출하세요.

[JSON 포맷]
{{
  "representative_position": "문자열(최대 8글자 단어, 없으면 빈문자열)",
  "careers": [
    {{ "company": "회사명(미상)", "title": "직급/직무", "start_date": "YYYY-MM", "end_date": "YYYY-MM" }}
  ]
}}

[이력서 전문 (15000자까지)]
{raw_text[:15000]}"""
    try:
        res = model.generate_content(prompt).text.strip()
        data = json.loads(res)
        car = json.dumps(data.get('careers', []), ensure_ascii=False)
        pos = data.get('representative_position', '')
        return db_id, name_kr, raw_text, car, pos
    except:
        return db_id, name_kr, raw_text, None, None

def run_physical_rescue():
    # Load available files from both dirs
    available_files = {}
    for d in [DIR_V8, DIR_RAW]:
        if not os.path.exists(d): continue
        for f in os.listdir(d):
            if f.endswith('.docx'):
                available_files[f] = os.path.join(d, f)

    conn = sqlite3.connect(r'c:\Users\cazam\Downloads\이력서자동분석검색시스템\candidates.db')
    c = conn.cursor()
    c.execute("SELECT id, name_kr FROM candidates WHERE careers_json IS NULL OR careers_json='[]' OR careers_json=''")
    ghosts = c.fetchall()

    tasks = []
    for db_id, name_kr in ghosts:
        if not name_kr: continue
        clean_name = name_kr.split('(')[0].strip() # Extract "홍길동" from "홍길동(신입)"
        if not clean_name: continue
        
        matched_file = None
        for file_name, file_path in available_files.items():
            if clean_name in file_name:
                matched_file = file_path
                break
        
        if matched_file:
            tasks.append((db_id, name_kr, matched_file))

    print(f"Found {len(tasks)} physical files for {len(ghosts)} remaining ghosts.")
    if not tasks: return

    success = 0
    with ThreadPoolExecutor(max_workers=5) as p:
        futures = {p.submit(process_file, t): t for t in tasks}
        for fut in tqdm(as_completed(futures), total=len(futures), desc="Physical File Recovery"):
            db_id, n_kr, raw_text, car, pos = fut.result()
            if car is not None and car != '[]' and len(car)>5:
                c.execute("UPDATE candidates SET raw_text=?, careers_json=? WHERE id=?", (raw_text, car, db_id))
                
                if pos and '(' not in n_kr:
                    clean_pos = str(pos).replace('(', '').replace(')', '').strip()
                    if clean_pos:
                        new_name = f"{n_kr}({clean_pos})"
                        c.execute("UPDATE candidates SET name_kr=? WHERE id=?", (new_name, db_id))
                success += 1
                if success % 10 == 0: conn.commit()
    conn.commit()
    conn.close()
    print(f"Physical file rescue completed! Successfully resurrected {success} candidates.")

if __name__ == '__main__':
    run_physical_rescue()
